"""
Smart MOP Generator — Unified Master v1.0
==========================================
Consolidates:
  · Smart MOP Generator v5   (Solution Doc → MOP, Ericsson navy/blue UI)
  · Compiled MOP Generator v6 (Activity MOP media/attachment/comment injection)

Unified capabilities:
  · Solution Document upload           — REQUIRED  (all 12 sections, any scenario)
  · Activity MOP upload                — OPTIONAL  (images, OLE attachments, comments extracted & injected)
  · Additional uploads (logs, images,
    screenshots, tables, flowcharts)   — OPTIONAL  (referenced via placeholders in SOP)
  · Multi-language input support       — output always in professional English
  · Truncation-aware, ZDR-strict, automation-ready output
  · Zero Data Retention: all processing in-memory, nothing written to disk
"""

import io
import re
import time
import zipfile
import os
from datetime import datetime
from pathlib import Path
from copy import deepcopy
from lxml import etree

import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Smart MOP Generator",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Session state init ──
for _key, _val in {
    "output_bytes":       b"",
    "activity_name":      "",
    "today_str":          "",
    "sections":           {},
    "filled":             0,
    "images_n":           0,
    "total_n":            0,
    "failed_media":       [],
    "injected_media":     0,
    "comments_injected":  0,
    "comments_failed":    [],
    "authenticated":      False,
    "auth_attempts":      0,
}.items():
    if _key not in st.session_state:
        st.session_state[_key] = _val

# ─────────────────────────────────────────────────────────────────
# PASSWORD GATE — shown before anything else
# ─────────────────────────────────────────────────────────────────
_CORRECT_PASSWORD = "Ericsson@1876"
_MAX_ATTEMPTS     = 5

def _render_login():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Lato:wght@300;400;700;900&display=swap');
    html, body, [class*="css"] { font-family:'Lato',sans-serif; background-color:#F4F6F9; }
    .login-wrap {
        max-width: 420px; margin: 6vh auto 0; padding: 2.8rem 2.4rem 2.2rem;
        background: #ffffff; border-radius: 16px;
        border: 1px solid #dde4ed; border-top: 4px solid #0082C8;
        box-shadow: 0 8px 32px rgba(0,51,102,0.10);
        text-align: center;
    }
    .login-logo  { font-family:'Lato',sans-serif; font-weight:900; font-size:1.8rem;
                   letter-spacing:5px; color:#0082C8; text-transform:uppercase; }
    .login-sub   { font-size:0.62rem; letter-spacing:2px; color:#9aaab8;
                   text-transform:uppercase; margin-top:3px; }
    .login-title { font-size:1.1rem; font-weight:700; color:#003366;
                   margin:1.4rem 0 0.3rem; }
    .login-desc  { font-size:0.76rem; color:#9aaab8; margin-bottom:1.6rem; }
    .login-err   { background:rgba(200,60,0,0.07); border:1px solid rgba(200,60,0,0.22);
                   border-left:4px solid #cc3300; border-radius:0 8px 8px 0;
                   padding:0.6rem 1rem; font-size:0.78rem; color:#7a2200;
                   margin-top:0.8rem; text-align:left; }
    .login-lock  { background:rgba(200,60,0,0.05); border:1px solid rgba(200,60,0,0.20);
                   border-radius:10px; padding:1.2rem; font-size:0.82rem;
                   color:#7a2200; font-weight:600; margin-top:0.8rem; }
    .login-zdr   { font-size:0.65rem; color:#b0bec5; margin-top:1.2rem;
                   border-top:1px solid #eee; padding-top:0.9rem; }
    </style>
    """, unsafe_allow_html=True)

    st.markdown('<div class="login-wrap">', unsafe_allow_html=True)
    st.markdown("""
        <div class="login-logo">ERICSSON</div>
        <div class="login-sub">Telecom Automation Toolkit</div>
        <div class="login-title">🔐 Smart MOP Generator</div>
        <div class="login-desc">Enter your access password to continue</div>
    """, unsafe_allow_html=True)

    attempts_left = _MAX_ATTEMPTS - st.session_state["auth_attempts"]

    if st.session_state["auth_attempts"] >= _MAX_ATTEMPTS:
        st.markdown("""
        <div class="login-lock">
            🔒 Access locked — maximum attempts reached.<br>
            Please refresh the page to try again.
        </div>
        """, unsafe_allow_html=True)
    else:
        pwd = st.text_input(
            "Password", type="password",
            placeholder="Enter password…",
            label_visibility="collapsed"
        )
        if st.button("🔓  Login", use_container_width=True):
            if pwd == _CORRECT_PASSWORD:
                st.session_state["authenticated"] = True
                st.rerun()
            else:
                st.session_state["auth_attempts"] += 1
                remaining = _MAX_ATTEMPTS - st.session_state["auth_attempts"]
                if remaining > 0:
                    st.markdown(
                        f'<div class="login-err">❌ Incorrect password — '
                        f'{remaining} attempt(s) remaining.</div>',
                        unsafe_allow_html=True
                    )
                else:
                    st.markdown(
                        '<div class="login-lock">🔒 Access locked — '
                        'maximum attempts reached. Please refresh to try again.</div>',
                        unsafe_allow_html=True
                    )

    st.markdown("""
        <div class="login-zdr">
            🔒 Zero Data Retention &nbsp;·&nbsp; All processing in-memory only
            &nbsp;·&nbsp; Ericsson Internal Tool
        </div>
    """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

if not st.session_state["authenticated"]:
    _render_login()
    st.stop()

# ─────────────────────────────────────────────────────────────────
# CSS — Ericsson Navy/Blue palette: #001f4d, #003366, #0082C8
# ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Lato:wght@300;400;700;900&family=Source+Code+Pro:wght@400;600&display=swap');

html, body, [class*="css"] {
  font-family: 'Lato', sans-serif;
  background-color: #F4F6F9;
  color: #1A1A2E;
}
.block-container { padding-top: 1.5rem !important; padding-bottom: 2rem; max-width: 100%; }

[data-testid="stSidebar"] {
  background: linear-gradient(180deg, #001020 0%, #001f4d 40%, #003366 100%) !important;
  border-right: 1px solid rgba(0,130,200,0.2);
}
[data-testid="stSidebar"] * { color: #e8f0ff !important; }
[data-testid="stSidebar"] .stMarkdown h1,
[data-testid="stSidebar"] .stMarkdown h2,
[data-testid="stSidebar"] .stMarkdown h3 { color: #ffffff !important; }
[data-testid="stSidebar"] hr { border-color: rgba(0,130,200,0.3) !important; }
[data-testid="stSidebar"] label { color: #90b8e0 !important; font-size: .78rem !important; }

.eri-topbar {
  background: linear-gradient(90deg, #001020, #001f4d, #003366);
  border-bottom: 3px solid #0082C8;
  padding: 1rem 2rem 0.8rem;
  border-radius: 12px;
  margin-bottom: 1.5rem;
  display: flex;
  align-items: center;
  justify-content: space-between;
}
.eri-logo-text { font-family:'Lato',sans-serif; font-weight:900; font-size:1.5rem; letter-spacing:3px; color:#0082C8; text-transform:uppercase; }
.eri-logo-sub  { font-size:0.7rem; letter-spacing:1.5px; color:rgba(255,255,255,0.4); text-transform:uppercase; }
.eri-app-title { font-size:1.15rem; font-weight:700; color:#ffffff; letter-spacing:0.3px; }
.eri-app-sub   { font-size:0.72rem; color:rgba(255,255,255,0.45); letter-spacing:0.5px; margin-top:2px; }
.eri-version   { background:rgba(0,130,200,0.15); border:1px solid rgba(0,130,200,0.3); border-radius:20px; padding:3px 12px; font-size:0.65rem; color:#0082C8; font-weight:700; letter-spacing:1px; }

.priv-bar {
  background: rgba(0,64,128,0.06);
  border: 1px solid rgba(0,130,200,0.18);
  border-left: 4px solid #0082C8;
  border-radius: 0 8px 8px 0;
  padding: 0.65rem 1rem;
  font-size: 0.78rem;
  color: #003366;
  margin-bottom: 1.2rem;
}
.priv-bar strong { color: #0082C8; }

.warn-bar {
  background: rgba(200,80,0,0.06);
  border: 1px solid rgba(200,80,0,0.22);
  border-left: 4px solid #cc5500;
  border-radius: 0 8px 8px 0;
  padding: 0.65rem 1rem;
  font-size: 0.78rem;
  color: #7a3000;
  margin-bottom: 0.8rem;
}
.warn-bar strong { color: #cc5500; }

.info-bar {
  background: rgba(0,64,200,0.05);
  border: 1px solid rgba(0,130,200,0.20);
  border-left: 4px solid #0082C8;
  border-radius: 0 8px 8px 0;
  padding: 0.65rem 1rem;
  font-size: 0.78rem;
  color: #003366;
  margin-bottom: 0.8rem;
}

.eri-card {
  background: #ffffff;
  border: 1px solid #dde4ed;
  border-radius: 12px;
  padding: 1.4rem 1.6rem;
  margin-bottom: 1rem;
  box-shadow: 0 2px 8px rgba(0,51,102,0.06);
  transition: box-shadow 0.2s, border-color 0.2s;
}
.eri-card:hover { border-color: #0082C8; box-shadow: 0 4px 16px rgba(0,130,200,0.10); }
.eri-card-title {
  font-size:0.68rem; font-weight:900; letter-spacing:1.8px; text-transform:uppercase;
  color:#003366; margin-bottom:0.9rem; display:flex; align-items:center; gap:8px;
}
.step-badge    { background:#003366; color:#ffffff; font-size:0.58rem; font-weight:700; padding:2px 8px; border-radius:4px; letter-spacing:0.5px; }
.optional-badge{ background:#0082C8; color:#ffffff; font-size:0.55rem; font-weight:700; padding:2px 7px; border-radius:4px; letter-spacing:0.4px; }

.pill-ok   { display:inline-flex; align-items:center; gap:6px; background:rgba(0,150,80,0.08); border:1px solid rgba(0,150,80,0.22); border-radius:6px; padding:5px 12px; font-size:0.76rem; color:#006633; margin:3px 0; }
.pill-warn { display:inline-flex; align-items:center; gap:6px; background:rgba(200,80,0,0.07); border:1px solid rgba(200,80,0,0.2); border-radius:6px; padding:5px 12px; font-size:0.76rem; color:#7a3500; margin:3px 0; }
.pill-info { display:inline-flex; align-items:center; gap:6px; background:rgba(0,82,130,0.07); border:1px solid rgba(0,130,200,0.2); border-radius:6px; padding:5px 12px; font-size:0.76rem; color:#003366; margin:3px 0; }

.stButton > button {
  background: linear-gradient(135deg,#003366,#0082C8) !important;
  color: #ffffff !important; border:none !important; border-radius:8px !important;
  font-family:'Lato',sans-serif !important; font-weight:700 !important;
  font-size:0.9rem !important; padding:0.6rem 2rem !important; width:100% !important;
  letter-spacing:0.4px !important; transition:all 0.2s !important;
}
.stButton > button:hover { background:linear-gradient(135deg,#004080,#0099E6) !important; transform:translateY(-1px) !important; box-shadow:0 6px 20px rgba(0,130,200,0.28) !important; }
.stButton > button:disabled { opacity:0.38 !important; transform:none !important; }
[data-testid="stDownloadButton"] > button {
  background: linear-gradient(135deg,#003366,#005999) !important;
  color:#ffffff !important; border:none !important; border-radius:8px !important;
  font-family:'Lato',sans-serif !important; font-weight:700 !important;
  font-size:0.9rem !important; padding:0.6rem 2rem !important; width:100% !important;
}
[data-testid="stDownloadButton"] > button:hover { background:linear-gradient(135deg,#004080,#0082C8) !important; transform:translateY(-1px) !important; }

.prog-wrap { background:#f8fafc; border:1px solid #dde4ed; border-radius:10px; padding:1rem 1.2rem; }
.ps { display:flex; align-items:center; gap:10px; padding:6px 0; font-size:0.78rem; border-bottom:1px solid rgba(0,0,0,0.04); }
.ps:last-child { border-bottom:none; }
.ps.done { color:#006633; font-weight:600; }
.ps.doing { color:#0082C8; font-weight:600; }
.ps.wait  { color:#9aaab8; }
.pd { width:8px; height:8px; border-radius:50%; flex-shrink:0; }
.pd.done  { background:#00a050; }
.pd.doing { background:#0082C8; animation:pulse 1s infinite; }
.pd.wait  { background:#c8d4e0; }
@keyframes pulse { 0%,100%{opacity:1;transform:scale(1)} 50%{opacity:.5;transform:scale(0.85)} }

.success-card {
  background:linear-gradient(135deg,rgba(0,51,102,0.04),rgba(0,130,200,0.06));
  border:1px solid rgba(0,130,200,0.25); border-top:4px solid #0082C8;
  border-radius:12px; padding:1.6rem; margin:0.8rem 0; text-align:center;
}
.success-icon  { font-size:2.2rem; margin-bottom:0.4rem; }
.success-title { font-size:1.1rem; font-weight:900; color:#003366; margin-bottom:0.25rem; }
.success-sub   { font-size:0.78rem; color:#0082C8; }
.success-name  { color:#003366; font-weight:700; }

.media-fail-card {
  background:rgba(200,60,0,0.04);
  border:1px solid rgba(200,60,0,0.20);
  border-left:4px solid #cc3300;
  border-radius:0 10px 10px 0;
  padding:1rem 1.2rem;
  margin-top:0.8rem;
}
.media-fail-title { font-size:0.75rem; font-weight:900; color:#cc3300; margin-bottom:0.5rem; letter-spacing:0.8px; text-transform:uppercase; }
.media-fail-item  { font-size:0.76rem; color:#7a2200; padding:3px 0; border-bottom:1px solid rgba(200,60,0,0.08); }
.media-fail-item:last-child { border-bottom:none; }

.metric-row { display:grid; grid-template-columns:repeat(4,1fr); gap:12px; margin-top:1rem; }
.metric-box { background:#f0f5fb; border:1px solid #dde4ed; border-radius:10px; padding:1rem; text-align:center; }
.metric-val { font-family:'Lato',sans-serif; font-size:1.7rem; font-weight:900; color:#003366; }
.metric-sub { font-size:0.58rem; color:#9aaab8; margin-top:1px; font-weight:700; letter-spacing:.6px; text-transform:uppercase; }

.sidebar-info { background:rgba(0,130,200,0.08); border:1px solid rgba(0,130,200,0.18); border-radius:8px; padding:0.8rem 1rem; font-size:0.74rem; color:#b8d4f0 !important; margin:0.5rem 0; }
hr { border-color:rgba(0,51,102,0.1) !important; }
.stFileUploader label { color:#003366 !important; font-weight:600 !important; font-size:0.82rem !important; }
.footer { text-align:center; font-size:0.65rem; color:#9aaab8; padding:1rem 0 0.4rem; border-top:1px solid #dde4ed; letter-spacing:0.5px; margin-top:1rem; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────
HEADING_MAP = {
    "objective":            ["objective"],
    "activity_description": ["activity description"],
    "activity_type":        ["activity type"],
    "domain_in_scope":      ["domain in scope", "domain"],
    "prerequisites":        ["pre-requisites", "prerequisites"],
    "inventory_details":    ["inventory details", "inventory"],
    "node_connectivity":    ["node connectivity", "node connectivity process"],
    "iam":                  ["identity & access", "identity and access", "identity"],
    "triggering_method":    ["activity triggering", "triggering method"],
    "sop":                  ["standard operating procedure", "sop"],
    "acceptance_criteria":  ["acceptance criteria"],
    "assumptions":          ["assumptions"],
    "connectivity_diagram": ["connectivity diagram"],
}

SECTION_KEYS = [
    "objective", "activity_description", "activity_type", "domain_in_scope",
    "prerequisites", "inventory_details", "node_connectivity", "iam",
    "triggering_method", "sop", "acceptance_criteria", "assumptions",
    "connectivity_diagram",
]

SECTION_LABELS = {
    "objective":            "1. Objective",
    "activity_description": "2. Activity Description",
    "activity_type":        "3. Activity Type",
    "domain_in_scope":      "4. Domain in Scope",
    "prerequisites":        "5. Pre-requisites",
    "inventory_details":    "6. Inventory Details (Node Name, Type, Count, Vendor)",
    "node_connectivity":    "7. Node Connectivity Process",
    "iam":                  "8. Identity and Access Management",
    "triggering_method":    "9. Activity Triggering Method",
    "sop":                  "10. Standard Operating Procedure",
    "acceptance_criteria":  "11. Acceptance Criteria (UAT Scenarios)",
    "assumptions":          "12. Assumptions",
    "connectivity_diagram": "Connectivity Diagram",
}

IMAGE_PLACEHOLDER_RE = re.compile(
    r'\[IMAGE[^\]]*\]|\[SCREENSHOT[^\]]*\]|\[ATTACHMENT[^\]]*\]'
    r'|\[IMAGE/SCREENSHOT[^\]]*\]|\[DIAGRAM[^\]]*\]|\[FIGURE[^\]]*\]',
    re.IGNORECASE
)

_NS_R    = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_NS_BLIP = "http://schemas.openxmlformats.org/drawingml/2006/main"
_NS_W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_V    = "urn:schemas-microsoft-com:vml"
_HEADING_COLOR_HEX = "1F497D"


# ─────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────
def normalize_heading(text: str):
    t = re.sub(r'^\d+[\.]\s*', '', text).strip().lower()
    t = re.sub(r'\s+', ' ', t)
    t = re.sub(r'^section\s+\d+\s*[—\-]\s*', '', t)
    for key, aliases in HEADING_MAP.items():
        for alias in aliases:
            if alias in t:
                return key
    return None


def discover_templates() -> list:
    tmpl_dir = Path("templates")
    if tmpl_dir.exists():
        return sorted(tmpl_dir.glob("*.docx"))
    return []


def load_template_bytes(path: Path) -> bytes:
    with open(path, "rb") as f:
        return f.read()


# ─────────────────────────────────────────────────────────────────
# ACTIVITY MOP — MEDIA EXTRACTOR
# ─────────────────────────────────────────────────────────────────
class MediaItem:
    def __init__(self, kind, blob, ext, rId, position_index,
                 context_text="", filename="", prog_id="", content_type=""):
        self.kind           = kind
        self.blob           = blob
        self.ext            = ext
        self.rId            = rId
        self.position_index = position_index
        self.context_text   = context_text
        self.filename       = filename
        self.prog_id        = prog_id
        self.content_type   = content_type
        self.injected       = False
        self.inject_error   = None

    @property
    def display_name(self):
        if self.filename:
            return self.filename
        if "Excel" in self.prog_id:
            return f"Excel_Attachment_{self.position_index+1}.xlsx"
        return f"Attachment_{self.position_index+1}.{self.ext}"


def _recover_filename_from_ole(blob: bytes) -> str:
    try:
        text = blob.decode("latin-1", errors="replace")
        for m in re.finditer(r'[\x20-\x7E]{3,200}', text):
            candidate = m.group(0).strip()
            if re.search(r'\.(txt|log|xlsx|xls|docx|doc|csv|pdf|zip|bin)$',
                         candidate, re.IGNORECASE):
                name = candidate.replace("\\", "/").split("/")[-1].strip()
                if 3 < len(name) < 120:
                    return name
    except Exception:
        pass
    return ""


def extract_media_from_activity_mop(mop_bytes: bytes) -> list:
    """
    Walk Activity MOP paragraphs in strict document order.
    Collect real PNG/JPG screenshots and OLE/Package attachments.
    Returns a flat list of MediaItem in document order.
    """
    doc = Document(io.BytesIO(mop_bytes))
    media_items: list = []
    position    = 0
    prev_text   = ""

    img_rel_map = {}
    att_rel_map = {}

    for rId, rel in doc.part.rels.items():
        rt = rel.reltype.split("/")[-1]
        try:
            if rt == "image":
                ct  = rel.target_part.content_type
                raw = ct.split("/")[-1].lower()
                ext = "jpg" if raw == "jpeg" else raw
                img_rel_map[rId] = (rel.target_part.blob, ext, ct)
            elif rt in ("oleObject", "package"):
                blob   = rel.target_part.blob
                target = rel.target_ref
                if rt == "package":
                    fname = target.split("/")[-1]
                    ext   = fname.split(".")[-1].lower() if "." in fname else "bin"
                else:
                    fname = _recover_filename_from_ole(blob)
                    ext   = fname.split(".")[-1].lower() if fname and "." in fname else "bin"
                att_rel_map[rId] = (blob, ext, fname)
        except Exception:
            pass

    for para in doc.paragraphs:
        text    = para.text.strip()
        p_xml   = para._p
        xml_str = etree.tostring(p_xml, encoding="unicode")

        ole_rids_here = re.findall(
            r'<[^>]*OLEObject[^>]+r:id="(rId\d+)"[^>]*/?>',
            xml_str
        )
        ole_rids_here += re.findall(
            r'<[^>]*OLEObject[^>]+r:id="(rId\d+)"',
            xml_str
        )
        ole_rids_here = list(dict.fromkeys(ole_rids_here))

        prog_map = {}
        for m in re.finditer(r'<[^>]*OLEObject[^>]+>', xml_str):
            tag   = m.group(0)
            rid_m = re.search(r'r:id="(rId\d+)"', tag)
            pro_m = re.search(r'ProgID="([^"]+)"', tag)
            if rid_m and pro_m:
                prog_map[rid_m.group(1)] = pro_m.group(1)

        icon_rids = set()
        for m in re.finditer(r'r:id="(rId\d+)"', xml_str):
            rid = m.group(1)
            if rid in img_rel_map and img_rel_map[rid][1].lower() in ("emf", "wmf"):
                icon_rids.add(rid)

        for ole_rid in ole_rids_here:
            if ole_rid not in att_rel_map:
                continue
            blob, ext, fname = att_rel_map[ole_rid]
            prog_id = prog_map.get(ole_rid, "")
            if "Excel" in prog_id or "Excel" in fname or ext == "xlsx":
                ext  = "xlsx"
                if not fname:
                    fname = f"Excel_Attachment_{position+1}.xlsx"
            item = MediaItem(
                kind="attachment", blob=blob, ext=ext,
                rId=ole_rid, position_index=position,
                context_text=prev_text, filename=fname, prog_id=prog_id
            )
            media_items.append(item)
            position += 1

        if not ole_rids_here:
            seen_in_para = set()
            for blip in p_xml.findall(f".//{{{_NS_BLIP}}}blip"):
                embed = blip.get(f"{{{_NS_R}}}embed")
                if not embed or embed not in img_rel_map or embed in seen_in_para:
                    continue
                blob, ext, ct = img_rel_map[embed]
                if ext.lower() in ("emf", "wmf"):
                    # Small EMF/WMF blobs are almost always decorative
                    # icons/logos (e.g. OLE preview icons) - safe to skip.
                    # Larger ones are likely a real pasted flowchart/diagram;
                    # python-docx can't render EMF/WMF directly, so we still
                    # let it flow through as a MediaItem - the existing
                    # insertion try/except will catch the failure and emit
                    # a "manual insertion needed" notice instead of the
                    # image silently disappearing with no trace.
                    if len(blob) < 3000:
                        continue
                elif len(blob) < 2000:
                    continue
                if embed in {m.rId for m in media_items}:
                    continue
                seen_in_para.add(embed)
                item = MediaItem(
                    kind="image", blob=blob, ext=ext,
                    rId=embed, position_index=position,
                    context_text=prev_text, content_type=ct
                )
                media_items.append(item)
                position += 1

        if text:
            prev_text = text

    return media_items


# ─────────────────────────────────────────────────────────────────
# ACTIVITY MOP — HEADING MAP & SECTION EXTRACTOR
# ─────────────────────────────────────────────────────────────────
# Maps Activity MOP heading keywords → our 12 section keys.
# Longer / more-specific phrases listed first (order matters for matching).
MOP_HEADING_MAP = [
    ("objective",                    "objective"),
    ("introduction",                 "objective"),
    ("overview",                     "objective"),
    ("purpose",                      "activity_description"),
    ("activity description",         "activity_description"),
    ("description",                  "activity_description"),
    ("background",                   "activity_description"),
    ("activity type",                "activity_type"),
    ("change type",                  "activity_type"),
    ("type of activity",             "activity_type"),
    ("domain in scope",              "domain_in_scope"),
    ("in scope",                     "domain_in_scope"),
    ("out of scope",                 "domain_in_scope"),
    ("scope",                        "domain_in_scope"),
    ("domain",                       "domain_in_scope"),
    ("pre-requisite",                "prerequisites"),
    ("prerequisite",                 "prerequisites"),
    ("precondition",                 "prerequisites"),
    ("requirement",                  "acceptance_criteria"),
    ("benefit",                      "objective"),
    ("inventory",                    "inventory_details"),
    ("node detail",                  "inventory_details"),
    ("equipment",                    "inventory_details"),
    ("network element",              "inventory_details"),
    ("node connectivity",            "node_connectivity"),
    ("connectivity process",         "node_connectivity"),
    ("connection detail",            "node_connectivity"),
    ("access method",                "node_connectivity"),
    ("identity and access",          "iam"),
    ("identity & access",            "iam"),
    ("access management",            "iam"),
    ("authoris",                     "iam"),
    ("authoriz",                     "iam"),
    ("credential",                   "iam"),
    ("triggering method",            "triggering_method"),
    ("activity trigger",             "triggering_method"),
    ("trigger",                      "triggering_method"),
    ("scheduling",                   "triggering_method"),
    ("standard operating procedure", "sop"),
    ("method of procedure",          "sop"),
    ("work instruction",             "sop"),
    ("acceptance criteria",          "acceptance_criteria"),
    ("uat",                          "acceptance_criteria"),
    ("user acceptance",              "acceptance_criteria"),
    ("sign-off",                     "acceptance_criteria"),
    ("sign off",                     "acceptance_criteria"),
    ("validation criteria",          "acceptance_criteria"),
    ("assumption",                   "assumptions"),
    ("constraint",                   "assumptions"),
    ("dependency",                   "assumptions"),
]

_BOILERPLATE_RE = re.compile(
    r'^(contents|table of contents|revision history|document control'
    r'|header|footer|page \d|confidential|ericsson network automation'
    r'|prepared by|approved by|document number|version no|method of procedure$)',
    re.IGNORECASE
)
_TOC_LINE_RE = re.compile(r'^\d+[\.\)]\s+\w.{0,60}Page\s+\d+', re.IGNORECASE)


def normalize_mop_heading(text: str):
    """Map Activity MOP heading text to one of our 12 section keys, or None."""
    t = re.sub(r'^\d[\d\.]*\s*', '', text).strip().lower()
    t = re.sub(r'\s+', ' ', t)
    for keyword, sec_key in MOP_HEADING_MAP:
        if keyword in t:
            return sec_key
    return None


def extract_mop_sections(mop_bytes: bytes) -> dict:
    """
    Parse Activity MOP and distribute content into our 12 section buckets.
    Rules:
    - Heading paragraphs → mapped to section key via normalize_mop_heading
    - Unmatched headings → content goes to 'sop' (catch-all)
    - Tables → each row converted to a styled paragraph element
    - All body text formatting is normalised by _clone_para at render time
    """
    doc     = Document(io.BytesIO(mop_bytes))
    buckets = {k: [] for k in SECTION_KEYS}
    current = "sop"
    _W      = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _is_mop_heading(para):
        if para.style.name.startswith("Heading"):
            return True
        txt = para.text.strip()
        if not txt or len(txt) > 100:
            return False
        is_bold    = any(r.bold for r in para.runs if r.text.strip())
        is_allcaps = txt == txt.upper() and len(txt) > 3
        return is_bold or is_allcaps

    body_elem = doc.element.body
    for child in body_elem:
        tag = child.tag.split("}")[-1]

        if tag == "p":
            para_obj = None
            for p in doc.paragraphs:
                if p._p is child:
                    para_obj = p; break
            if para_obj is None:
                continue
            text = para_obj.text.strip()

            if _is_mop_heading(para_obj):
                mapped  = normalize_mop_heading(text) if text else None
                current = mapped if mapped else "sop"
                continue  # heading itself not copied

            if not text:
                # Skip empty pBdr paragraphs (separator lines)
                pBdr_chk = child.find(f".//{{{_W}}}pBdr")
                if pBdr_chk is not None:
                    continue
                buckets[current].append(deepcopy(child)); continue
            if _BOILERPLATE_RE.match(text): continue
            if _TOC_LINE_RE.match(text):    continue

            buckets[current].append(deepcopy(child))

        elif tag == "tbl":
            # Keep table as actual tbl element — will be cloned by _clone_table at render time
            buckets[current].append(deepcopy(child))

    # Strip trailing empty paragraphs
    for key in buckets:
        while buckets[key]:
            last_txt = "".join(
                t.text or ""
                for t in buckets[key][-1].findall(".//" + qn("w:t"))
            ).strip()
            if not last_txt:
                buckets[key].pop()
            else:
                break

    return buckets



def extract_activity_name(doc: Document) -> str:
    paragraphs = doc.paragraphs
    obj_idx = None
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading") and normalize_heading(text) == "objective":
            obj_idx = i
            break
        if normalize_heading(text) == "objective":
            obj_idx = i
            break

    if obj_idx is not None and obj_idx > 0:
        for i in range(obj_idx - 1, -1, -1):
            text = paragraphs[i].text.strip()
            if not text:
                continue
            upper = text.upper()
            if upper in ("METHOD OF PROCEDURE", "METHOD OF PROCEDURE (MOP)",
                         "CONTENTS:", "CONTENTS", "TITLE PAGE"):
                continue
            if re.match(r'^\d+[\.]\s+\w.*Page\s+\d+', text):
                continue
            if "\n" in text or re.match(
                    r'^(Customer|Activity Title|Document Reference|Domain|Vendor)[\s]*:',
                    text, re.IGNORECASE):
                for line in text.split("\n"):
                    line = line.strip()
                    m = re.match(r'^Activity\s+Title\s*:\s*(.+)', line, re.IGNORECASE)
                    if m:
                        return m.group(1).strip()
                continue
            if normalize_heading(text) is not None:
                continue
            if re.match(r'^(Customer|Header|Footer|Document)[\s]*:', text, re.IGNORECASE):
                continue
            name = re.sub(r'^MOP\s*:\s*', '', text, flags=re.IGNORECASE)
            name = re.sub(r'^UC\s*:\s*', '', name, flags=re.IGNORECASE)
            name = re.sub(r'^Activity\s+Title\s*:\s*', '', name, flags=re.IGNORECASE)
            name = re.sub(r'^Method of Procedure\s*[\(\[]?MOP[\)\]]?\s*[:\-]?\s*',
                          '', name, flags=re.IGNORECASE)
            name = name.strip()
            if name and len(name) > 3:
                return name

    for para in paragraphs[:10]:
        if para.style.name.startswith("Heading 1"):
            name = para.text.strip()
            name = re.sub(r'^MOP\s*:\s*', '', name, flags=re.IGNORECASE)
            if name and normalize_heading(name) is None:
                return name

    for para in paragraphs[:15]:
        for run in para.runs:
            if run.italic and run.underline and para.text.strip():
                return para.text.strip()

    return "Activity Name"


def extract_sections(doc: Document) -> dict:
    """
    Extract solution document section content as XML element lists.
    Preserves tables, empty paragraphs, and ALL run formatting exactly as-is
    (solution doc formatting is the source of truth — it goes through _clone_para
    only to strip page-break artefacts, not to change colors/bold/font).
    """
    sections    = {k: [] for k in SECTION_KEYS}
    current_key = None
    _W          = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    body_elem = doc.element.body
    for child in body_elem:
        tag = child.tag.split("}")[-1]
        if tag == "sectPr":
            continue

        if tag == "tbl":
            # Table: clone verbatim, assign to current section
            if current_key and current_key in sections:
                sections[current_key].append(deepcopy(child))
            continue

        if tag != "p":
            continue

        para_obj = None
        for p in doc.paragraphs:
            if p._p is child:
                para_obj = p; break

        style = para_obj.style.name if para_obj else ""
        text  = para_obj.text.strip() if para_obj else (
            "".join(t.text or "" for t in child.findall(".//" + qn("w:t"))).strip()
        )

        is_h1            = (style.startswith("Heading 1") or style == "Heading1")
        is_h2            = (style.startswith("Heading 2") or style == "Heading2")
        is_heading_style = style.startswith("Heading")
        key_from_text    = normalize_heading(text) if text else None

        # ── Heading 1 and Heading 2 are ALWAYS section boundaries ─────────────
        # Solution documents generated by Copilot use Heading 2 for all 12
        # section titles (SECTION 1 — OBJECTIVE, etc.) and Heading 1 only for
        # the top-level document title.  Both must be treated as boundaries so
        # that current_key advances correctly and content fills the right bucket.
        # Heading 3 and below (e.g. "Monitoring Phase", "Tracing Phase" inside
        # SOP) all return None from normalize_heading() so they never cause a
        # wrong section switch — they fall through as body content instead.
        if is_h1 or is_h2:
            if key_from_text:
                current_key = key_from_text
            # Always skip — section headings are never body content
            continue

        # ── Heading 3+ : NEVER switch sections — always body content ──────────
        # Heading3 paragraphs (e.g. "Monitoring Phase", "Tracing Phase" under
        # SOP) must NOT trigger a section switch even if their text contains a
        # keyword.  Only Heading1/2 define section boundaries.
        # They fall through and are added as body content under current_key.
        if is_heading_style:
            pass  # fall through to body-content handling

        # ── Plain text paragraphs NEVER switch sections ─────────────────────────
        # Only Heading1/2 define section boundaries. A plain paragraph containing
        # a keyword like "inventory" (e.g. "Identify vendor type using inventory
        # lookup.") must NOT trigger a section switch.
        # This block intentionally does nothing — falls through to body content.

        if current_key is None:
            continue

        # Filter boilerplate
        if text.upper() in ("METHOD OF PROCEDURE", "METHOD OF PROCEDURE (MOP)",
                            "CONTENTS:", "CONTENTS"):
            continue
        if re.match(r'^\d+\.\s+\w.*Page\s+\d+', text):
            continue
        if re.match(r'^(Customer|Header|Footer|Activity Title|Document)[\s]*:',
                    text, re.IGNORECASE):
            continue
        if text == "sample...":
            continue

        # Skip empty paragraphs that only have a pBdr (horizontal separator lines)
        # These come from the source doc and would render as dark lines in output
        if not text:
            pBdr_check = child.find(f".//{{{_W}}}pBdr")
            if pBdr_check is not None:
                continue   # drop — it's just a separator line

        # Include para — even empty ones (carry spacing/formatting)
        if current_key in sections:
            sections[current_key].append(deepcopy(child))

    return sections


# ─────────────────────────────────────────────────────────────────
# COMMENT EXTRACTOR FROM ACTIVITY MOP
# ─────────────────────────────────────────────────────────────────
class CommentItem:
    def __init__(self, cid, author, date, text, para_text, para_index):
        self.cid        = cid
        self.author     = author
        self.date       = date
        self.text       = text
        self.para_text  = para_text
        self.para_index = para_index


def extract_comments_from_mop(mop_bytes: bytes) -> tuple:
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(io.BytesIO(mop_bytes)) as z:
        names = z.namelist()
        def _read(path):
            return z.read(path).decode("utf-8") if path in names else ""
        comments_xml     = _read("word/comments.xml")
        comments_ext_xml = _read("word/commentsExtended.xml")
        comments_ids_xml = _read("word/commentsIds.xml")
        comments_exs_xml = _read("word/commentsExtensible.xml")
        doc_xml          = _read("word/document.xml")

    if not comments_xml:
        return [], "", "", "", ""

    comments_root = etree.fromstring(comments_xml.encode())
    comment_map   = {}
    for c in comments_root.findall(f"{{{_W}}}comment"):
        cid    = c.get(f"{{{_W}}}id")
        author = c.get(f"{{{_W}}}author", "Unknown")
        date   = c.get(f"{{{_W}}}date", "")
        texts  = [t.text or "" for t in c.findall(f".//{{{_W}}}t")]
        text   = "".join(texts).strip()
        comment_map[cid] = CommentItem(cid, author, date, text, "", -1)

    doc_root = etree.fromstring(doc_xml.encode())
    body     = doc_root.find(f"{{{_W}}}body")
    paras    = body.findall(f".//{{{_W}}}p")

    for i, para in enumerate(paras):
        para_texts = [t.text or "" for t in para.findall(f".//{{{_W}}}t")]
        para_text  = "".join(para_texts).strip()
        starts = para.findall(f"{{{_W}}}commentRangeStart")
        refs   = para.findall(f".//{{{_W}}}commentReference")
        anchored_ids = set()
        for s in starts:
            anchored_ids.add(s.get(f"{{{_W}}}id"))
        for r in refs:
            anchored_ids.add(r.get(f"{{{_W}}}id"))
        for cid in anchored_ids:
            if cid in comment_map:
                comment_map[cid].para_text  = para_text
                comment_map[cid].para_index = i

    items = sorted(comment_map.values(), key=lambda x: x.para_index)
    return items, comments_xml, comments_ext_xml, comments_ids_xml, comments_exs_xml


def _inject_comments_into_docx(doc_bytes, comment_items, comments_xml,
                                comments_ext, comments_ids, comments_exs) -> tuple:
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    if not comment_items or not comments_xml:
        return doc_bytes, [], []

    in_buf  = io.BytesIO(doc_bytes)
    out_buf = io.BytesIO()
    with zipfile.ZipFile(in_buf, "r") as zin:
        doc_xml_bytes  = zin.read("word/document.xml")
        rels_xml_bytes = zin.read("word/_rels/document.xml.rels")
        ct_xml_bytes   = zin.read("[Content_Types].xml")

    doc_root  = etree.fromstring(doc_xml_bytes)
    body      = doc_root.find(f"{{{_W}}}body")
    out_paras = body.findall(f".//{{{_W}}}p")

    injected_ids = []
    failed_ids   = []

    for ci in comment_items:
        cid = ci.cid
        matched_para = None
        best_score   = 0
        if ci.para_text:
            for para in out_paras:
                texts = [t.text or "" for t in para.findall(f".//{{{_W}}}t")]
                ptxt  = "".join(texts).strip()
                if ptxt and ci.para_text:
                    src_words = set(ci.para_text.lower().split())
                    tgt_words = set(ptxt.lower().split())
                    common    = len(src_words & tgt_words)
                    score     = common / max(len(src_words), 1)
                    if score > best_score and score >= 0.5:
                        best_score   = score
                        matched_para = para

        if matched_para is None and out_paras:
            ratio      = ci.para_index / max(len(out_paras), 1)
            fallback_i = min(int(ratio * len(out_paras)), len(out_paras) - 1)
            matched_para = out_paras[fallback_i]

        if matched_para is None:
            failed_ids.append(cid)
            continue

        def _make_range_start(cv):
            el = OxmlElement("w:commentRangeStart")
            el.set(f"{{{_W}}}id", cv)
            return el
        def _make_range_end(cv):
            el = OxmlElement("w:commentRangeEnd")
            el.set(f"{{{_W}}}id", cv)
            return el
        def _make_reference(cv):
            r   = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            rs  = OxmlElement("w:rStyle")
            rs.set(f"{{{_W}}}val", "CommentReference")
            rPr.append(rs)
            r.append(rPr)
            cr = OxmlElement("w:commentReference")
            cr.set(f"{{{_W}}}id", cv)
            r.append(cr)
            return r

        pPr        = matched_para.find(f"{{{_W}}}pPr")
        children   = list(matched_para)
        insert_pos = (children.index(pPr) + 1) if pPr is not None else 0
        matched_para.insert(insert_pos, _make_range_start(cid))
        matched_para.append(_make_range_end(cid))
        matched_para.append(_make_reference(cid))
        injected_ids.append(cid)

    # Use the original document.xml bytes as base, then patch in comment anchors
    # by serialising only the modified body back. This preserves all namespace
    # declarations that Word requires and avoids "unreadable content" errors.
    new_doc_xml = etree.tostring(doc_root, xml_declaration=True,
                                 encoding="UTF-8", standalone=True,
                                 with_tail=True)
    rels_xml = rels_xml_bytes.decode("utf-8")

    COMMENT_REL     = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
    COMMENT_EXT_REL = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
    COMMENT_IDS_REL = "http://schemas.microsoft.com/office/2016/09/relationships/commentsIds"
    COMMENT_EXS_REL = "http://schemas.microsoft.com/office/2020/relationships/commentsExtensible"

    def _add_rel(rels, rid, rtype, target):
        # Skip if the relationship type OR target already exists (prevents duplicate/conflicting entries)
        if target in rels or rtype in rels:
            return rels
        new = (f'<Relationship Id="{rid}" Type="{rtype}" Target="{target}"/>')
        rels = rels.replace("</Relationships>", new + "</Relationships>")
        return rels

    rels_xml = _add_rel(rels_xml, "rIdCom1", COMMENT_REL,     "comments.xml")
    rels_xml = _add_rel(rels_xml, "rIdCom2", COMMENT_EXT_REL, "commentsExtended.xml")
    rels_xml = _add_rel(rels_xml, "rIdCom3", COMMENT_IDS_REL, "commentsIds.xml")
    rels_xml = _add_rel(rels_xml, "rIdCom4", COMMENT_EXS_REL, "commentsExtensible.xml")

    ct_xml = ct_xml_bytes.decode("utf-8")
    ct_entries = [
        ('word/comments.xml',
         'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'),
        ('word/commentsExtended.xml',
         'application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml'),
        ('word/commentsIds.xml',
         'application/vnd.openxmlformats-officedocument.wordprocessingml.commentsIds+xml'),
        ('word/commentsExtensible.xml',
         'application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtensible+xml'),
    ]
    for part_name, ct in ct_entries:
        if part_name not in ct_xml:
            override = f'<Override PartName="/{part_name}" ContentType="{ct}"/>'
            ct_xml = ct_xml.replace("</Types>", override + "</Types>")

    in_buf.seek(0)
    with zipfile.ZipFile(in_buf, "r") as zin, \
         zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        skip = {"word/document.xml", "word/_rels/document.xml.rels",
                "[Content_Types].xml", "word/comments.xml",
                "word/commentsExtended.xml", "word/commentsIds.xml",
                "word/commentsExtensible.xml"}
        for item in zin.infolist():
            if item.filename not in skip:
                # Preserve original compression type to avoid ZIP structure errors
                data = zin.read(item.filename)
                info = zipfile.ZipInfo(item.filename)
                info.compress_type = zipfile.ZIP_DEFLATED
                zout.writestr(info, data)
        zout.writestr("word/document.xml",             new_doc_xml)
        zout.writestr("word/_rels/document.xml.rels",  rels_xml.encode("utf-8"))
        zout.writestr("[Content_Types].xml",           ct_xml.encode("utf-8"))
        zout.writestr("word/comments.xml",             comments_xml.encode("utf-8"))
        if comments_ext:
            zout.writestr("word/commentsExtended.xml", comments_ext.encode("utf-8"))
        if comments_ids:
            zout.writestr("word/commentsIds.xml",      comments_ids.encode("utf-8"))
        if comments_exs:
            zout.writestr("word/commentsExtensible.xml", comments_exs.encode("utf-8"))

    out_buf.seek(0)
    return out_buf.read(), injected_ids, failed_ids


# ─────────────────────────────────────────────────────────────────
# DOCX BUILDER — CORE UTILITIES
# ─────────────────────────────────────────────────────────────────
def _apply_heading_color(p_elem):
    def _fix_color(rpr):
        color_el = rpr.find(qn("w:color"))
        if color_el is None:
            color_el = OxmlElement("w:color")
            rpr.append(color_el)
        color_el.set(qn("w:val"), _HEADING_COLOR_HEX)
        for attr in (qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade")):
            if color_el.get(attr) is not None:
                del color_el.attrib[attr]
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        p_rpr = pPr.find(qn("w:rPr"))
        if p_rpr is not None:
            _fix_color(p_rpr)
    for r_el in p_elem.findall(".//" + qn("w:r")):
        rpr = r_el.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            r_el.insert(0, rpr)
        if rpr.find(qn("w:b")) is None:
            rpr.insert(0, OxmlElement("w:b"))
        _fix_color(rpr)



def _ensure_mop_numbering(doc: Document) -> int:
    """
    Ensure the template document has a numbering definition where:
      ilvl=0 → decimal  (1, 2, 3)
      ilvl=1 → lowerLetter (a, b, c)
    Returns the numId to use for all list paragraphs.

    Strategy:
    1. Scan existing abstractNums for decimal+lowerLetter multilevel.
    2. If found, return its corresponding numId (lowest one that maps to it).
    3. If not found, inject a minimal abstractNum + num entry and return new numId.
    """
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Access numbering part
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    num_part = None
    try:
        num_part = doc.part.numbering_part
    except Exception:
        pass

    if num_part is None:
        # No numbering in template — inject a minimal one
        return _inject_minimal_numbering(doc)

    root = num_part._element

    # Find abstractNum with ilvl0=decimal, ilvl1=lowerLetter
    target_abstract_id = None
    for abst in root.findall(f"{{{_W}}}abstractNum"):
        lvls = abst.findall(f"{{{_W}}}lvl")
        if len(lvls) < 2:
            continue
        fmt0 = lvls[0].find(f"{{{_W}}}numFmt")
        fmt1 = lvls[1].find(f"{{{_W}}}numFmt")
        if fmt0 is None or fmt1 is None:
            continue
        v0 = fmt0.get(f"{{{_W}}}val", "")
        v1 = fmt1.get(f"{{{_W}}}val", "")
        if v0 == "decimal" and v1 == "lowerLetter":
            target_abstract_id = abst.get(f"{{{_W}}}abstractNumId")
            break

    if target_abstract_id is None:
        return _inject_minimal_numbering(doc)

    # Find the lowest numId that maps to this abstractNumId
    for num in root.findall(f"{{{_W}}}num"):
        ref = num.find(f"{{{_W}}}abstractNumId")
        if ref is not None and ref.get(f"{{{_W}}}val") == target_abstract_id:
            return int(num.get(f"{{{_W}}}numId", "1"))

    return _inject_minimal_numbering(doc)


def _inject_minimal_numbering(doc: Document) -> int:
    """
    Inject a minimal 2-level numbering definition into the document.
    Returns the new numId (9999).
    """
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NEW_ABSTRACT = "9998"
    NEW_NUM_ID   = "9999"

    abstract_xml = f"""<w:abstractNum xmlns:w="{_W}" w:abstractNumId="{NEW_ABSTRACT}">
  <w:multiLevelType w:val="multilevel"/>
  <w:lvl w:ilvl="0">
    <w:start w:val="1"/>
    <w:numFmt w:val="decimal"/>
    <w:lvlText w:val="%1."/>
    <w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="360" w:hanging="360"/></w:pPr>
    <w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="20"/></w:rPr>
  </w:lvl>
  <w:lvl w:ilvl="1">
    <w:start w:val="1"/>
    <w:numFmt w:val="lowerLetter"/>
    <w:lvlText w:val="%2."/>
    <w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>
    <w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="20"/></w:rPr>
  </w:lvl>
</w:abstractNum>"""

    num_xml = f"""<w:num xmlns:w="{_W}" w:numId="{NEW_NUM_ID}">
  <w:abstractNumId w:val="{NEW_ABSTRACT}"/>
</w:num>"""

    try:
        num_part = doc.part.numbering_part
        root     = num_part._element
        from lxml import etree as _et
        abst_el = _et.fromstring(abstract_xml)
        num_el  = _et.fromstring(num_xml)
        # Insert abstractNum before first <w:num>
        first_num = root.find(f"{{{_W}}}num")
        if first_num is not None:
            root.insert(list(root).index(first_num), abst_el)
        else:
            root.append(abst_el)
        root.append(num_el)
    except Exception:
        pass

    return int(NEW_NUM_ID)



    """Update {{current date}} placeholder in header AND cover page Date table cell.
    Footer/header tables (repeated on every page) are skipped by checking
    that the table's parent element is the document body directly.
    """
    import re as _re
    date_pat = _re.compile(r'^\d{4}-\d{2}-\d{2}$|^\d{2}-\d{2}-\d{4}$')
    body = doc.element.body
    # 1. Header placeholder
    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:
                if "{{current date}}" in run.text:
                    run.text = run.text.replace("{{current date}}", today_str)
    # 2. Cover page table Date cell — body tables only, first match only
    cover_done = False
    for table in doc.tables:
        if cover_done:
            break
        # Skip tables not directly in body (footer/header tables)
        if table._tbl.getparent() is not body:
            continue
        flat = [c.text.strip().lower() for c in table.rows[0].cells]
        if "version no." in flat or "version no" in flat:
            continue  # skip revision table
        # Look for a "Date" label in this table
        found_date_label = any("date" in h for h in flat)
        if not found_date_label:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip().lower() == "date":
                        found_date_label = True
                        break
                if found_date_label:
                    break
        if not found_date_label:
            continue
        # Replace date-value cell with today_str
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if date_pat.match(txt):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = today_str
                        else:
                            para.add_run(today_str).font.name = "Calibri"
                    cover_done = True


def _update_header_date(doc: Document, today_str: str):
    """Update {{current date}} placeholder in header AND cover page Date table cell."""
    import re as _re
    date_pat = _re.compile(r'^\d{4}-\d{2}-\d{2}$|^\d{2}-\d{2}-\d{4}$')
    body = doc.element.body
    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:
                if "{{current date}}" in run.text:
                    run.text = run.text.replace("{{current date}}", today_str)
    cover_done = False
    for table in doc.tables:
        if cover_done:
            break
        if table._tbl.getparent() is not body:
            continue
        flat = [c.text.strip().lower() for c in table.rows[0].cells]
        if "version no." in flat or "version no" in flat:
            continue
        found_date_label = any("date" in h for h in flat)
        if not found_date_label:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip().lower() == "date":
                        found_date_label = True
                        break
                if found_date_label:
                    break
        if not found_date_label:
            continue
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if date_pat.match(txt):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = today_str
                        else:
                            para.add_run(today_str).font.name = "Calibri"
                    cover_done = True



    for table in doc.tables:
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        if "Version No." not in header_cells:
            continue
        if len(table.rows) >= 2:
            row = table.rows[1]
            for para in row.cells[1].paragraphs:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = today_str
                else:
                    para.add_run(today_str).font.name = "Calibri"
            desc_idx = min(3, len(row.cells) - 1)
            for para in row.cells[desc_idx].paragraphs:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = activity_name
                else:
                    para.add_run(activity_name).font.name = "Calibri"
        break


def _update_revision_table(doc: Document, activity_name: str, today_str: str):
    for table in doc.tables:
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        if "Version No." not in header_cells:
            continue
        if len(table.rows) >= 2:
            row = table.rows[1]
            for para in row.cells[1].paragraphs:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = today_str
                else:
                    para.add_run(today_str).font.name = "Calibri"
            desc_idx = min(3, len(row.cells) - 1)
            for para in row.cells[desc_idx].paragraphs:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = activity_name
                else:
                    para.add_run(activity_name).font.name = "Calibri"
        break


def _make_xml_para(doc, text: str, bold=False, color_rgb=None,
                   italic=False, size_pt=11) -> etree._Element:
    p   = OxmlElement("w:p")
    r   = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fn  = OxmlElement("w:rFonts")
    fn.set(qn("w:ascii"), "Calibri"); fn.set(qn("w:hAnsi"), "Calibri")
    rpr.append(fn)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size_pt * 2)))
    rpr.append(sz)
    if bold:   rpr.append(OxmlElement("w:b"))
    if italic: rpr.append(OxmlElement("w:i"))
    if color_rgb:
        col = OxmlElement("w:color"); col.set(qn("w:val"), color_rgb)
        rpr.append(col)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t); p.append(r)
    return p



    tmp_p = doc.add_paragraph()
    run   = tmp_p.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Inches(width_inches))
    p_xml = tmp_p._element
    p_xml.getparent().remove(p_xml)
    return p_xml


def _make_notice_xml(desc: str) -> etree._Element:
    """
    Create a placeholder paragraph for media that could not be embedded.
    Shows a clearly labelled box so the engineer knows what to insert and where.
    """
    p   = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "FFF2CC")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4");   el.set(qn("w:color"), "CC3300")
        pBdr.append(el)
    pPr.append(pBdr)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:before"), "80"); sp.set(qn("w:after"), "80")
    pPr.append(sp)
    p.append(pPr)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fn = OxmlElement("w:rFonts"); fn.set(qn("w:ascii"), "Calibri"); fn.set(qn("w:hAnsi"), "Calibri")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "20")
    b  = OxmlElement("w:b")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "CC3300")
    rpr.append(fn); rpr.append(sz); rpr.append(b); rpr.append(col)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = f"[ ATTACHMENT / IMAGE REQUIRED: {desc} — Please insert manually here ]"
    r.append(t); p.append(r)
    return p


def _make_caption_xml() -> etree._Element:
    return _make_xml_para(
        None,
        "[Screenshot/Image — copied from Activity MOP]",
        italic=True, size_pt=9, color_rgb="595959"
    )


def _insert_after(anchor: etree._Element, new_elem: etree._Element):
    parent = anchor.getparent()
    idx    = list(parent).index(anchor)
    parent.insert(idx + 1, new_elem)


def _fitted_width_inches(img_bytes: bytes, max_w: float = 5.5, max_h: float = 8.0) -> float:
    """
    Pick an insertion width (inches) so the image fits within max_w x max_h,
    preserving its native aspect ratio. Prevents tall/portrait images
    (e.g. vertical flowcharts) from overflowing past a single page when
    forced to a fixed width. Falls back to max_w if dimensions can't be read.
    """
    try:
        from docx.image.image import Image as _DocxImage
        im = _DocxImage.from_blob(img_bytes)
        if not im.px_width or not im.px_height:
            return max_w
        aspect = im.px_width / im.px_height  # width / height
        width_in = min(max_w, max_h * aspect)
        return max(width_in, 1.0)
    except Exception:
        return max_w


def _make_image_xml(doc: Document, img_bytes: bytes, width_inches=None,
                     max_w: float = 5.5, max_h: float = 6.5):
    if width_inches is None:
        width_inches = _fitted_width_inches(img_bytes, max_w=max_w, max_h=max_h)
    tmp_p = doc.add_paragraph()
    run   = tmp_p.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Inches(width_inches))
    p_xml = tmp_p._element
    p_xml.getparent().remove(p_xml)
    return p_xml


def _new_section_num_id(doc: Document, abstract_id: str) -> str:
    """
    Create a FRESH w:num entry pointing to abstract_id so that numbering
    always restarts at 1 for each section.  Returns the new numId string.
    """
    from lxml import etree as _et
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        num_part = doc.part.numbering_part
    except Exception:
        return "2"
    root = num_part._element
    # Find highest existing numId
    existing = [
        int(n.get(f"{{{_W}}}numId", "0"))
        for n in root.findall(f"{{{_W}}}num")
    ]
    new_id = str(max(existing, default=0) + 1)
    num_el = _et.fromstring(
        f'<w:num xmlns:w="{_W}" w:numId="{new_id}">'
        f'<w:abstractNumId w:val="{abstract_id}"/>'
        f'</w:num>'
    )
    root.append(num_el)
    return new_id


def _get_decimal_abstract_id(doc: Document) -> str:
    """Return the abstractNumId for decimal/lowerLetter multilevel, or inject one."""
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        root = doc.part.numbering_part._element
    except Exception:
        return "2"
    for a in root.findall(f"{{{_W}}}abstractNum"):
        lvls = a.findall(f"{{{_W}}}lvl")
        if len(lvls) < 2:
            continue
        f0 = lvls[0].find(f"{{{_W}}}numFmt")
        f1 = lvls[1].find(f"{{{_W}}}numFmt")
        if f0 is None or f1 is None:
            continue
        if (f0.get(f"{{{_W}}}val") == "decimal" and
                f1.get(f"{{{_W}}}val") == "lowerLetter"):
            return a.get(f"{{{_W}}}abstractNumId")
    # Inject minimal abstractNum
    from lxml import etree as _et
    existing = [
        int(a.get(f"{{{_W}}}abstractNumId", "0"))
        for a in root.findall(f"{{{_W}}}abstractNum")
    ]
    new_aid = str(max(existing, default=0) + 1)
    abst = _et.fromstring(
        f'<w:abstractNum xmlns:w="{_W}" w:abstractNumId="{new_aid}">'
        f'<w:multiLevelType w:val="multilevel"/>'
        f'<w:lvl w:ilvl="0"><w:start w:val="1"/>'
        f'<w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/>'
        f'<w:lvlJc w:val="left"/>'
        f'<w:pPr><w:ind w:left="360" w:hanging="360"/></w:pPr>'
        f'<w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
        f'<w:sz w:val="20"/></w:rPr></w:lvl>'
        f'<w:lvl w:ilvl="1"><w:start w:val="1"/>'
        f'<w:numFmt w:val="lowerLetter"/><w:lvlText w:val="%2."/>'
        f'<w:lvlJc w:val="left"/>'
        f'<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
        f'<w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
        f'<w:sz w:val="20"/></w:rPr></w:lvl>'
        f'</w:abstractNum>'
    )
    first_num = root.find(f"{{{_W}}}num")
    if first_num is not None:
        root.insert(list(root).index(first_num), abst)
    else:
        root.append(abst)
    return new_aid




# Corporate body font — set explicitly on every body run
_BODY_FONT = "Ericsson Hilda Body"
_BODY_SIZE = "22"   # 11pt = 22 half-points


def _normalise_run(rpr, is_heading: bool = False):
    """
    Apply MOP house style to one rPr element in-place:
      • Font  → Ericsson Hilda Body (all 4 font slots)
      • Size  → 10pt
      • Color → 000000 black — strips any source color
      • Italic / underline / highlight / shd → stripped
    Headings are skipped (handled by _apply_heading_color).
    """
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    if is_heading:
        return
    for fn in rpr.findall(f"{{{_W}}}rFonts"):
        rpr.remove(fn)
    fn_new = OxmlElement("w:rFonts")
    fn_new.set(qn("w:ascii"),    _BODY_FONT)
    fn_new.set(qn("w:hAnsi"),    _BODY_FONT)
    fn_new.set(qn("w:eastAsia"), _BODY_FONT)
    fn_new.set(qn("w:cs"),       _BODY_FONT)
    rpr.insert(0, fn_new)
    for tag in (f"{{{_W}}}sz", f"{{{_W}}}szCs"):
        for el in rpr.findall(tag): rpr.remove(el)
    rpr.append(OxmlElement("w:sz"));   rpr[-1].set(qn("w:val"), _BODY_SIZE)
    rpr.append(OxmlElement("w:szCs")); rpr[-1].set(qn("w:val"), _BODY_SIZE)
    for c in rpr.findall(f"{{{_W}}}color"): rpr.remove(c)
    col = OxmlElement("w:color"); col.set(qn("w:val"), "000000"); rpr.append(col)
    for tag in (f"{{{_W}}}i", f"{{{_W}}}iCs", f"{{{_W}}}u",
                f"{{{_W}}}highlight", f"{{{_W}}}shd"):
        for el in rpr.findall(tag): rpr.remove(el)


def _import_numbering_defs(src_bytes: bytes, template_doc: Document) -> dict:
    """
    Import abstractNum + num definitions from src_bytes into template_doc,
    remapping IDs to avoid conflicts.
    Returns {original_numId_str: new_numId_str}.
    """
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    from lxml import etree as _et
    import zipfile as _zf

    try:
        with _zf.ZipFile(io.BytesIO(src_bytes)) as z:
            if 'word/numbering.xml' not in z.namelist():
                return {}
            src_root = _et.fromstring(z.read('word/numbering.xml'))
    except Exception:
        return {}

    try:
        tgt_root = template_doc.part.numbering_part._element
    except Exception:
        return {}

    existing_abs = [int(a.get(f"{{{_W}}}abstractNumId", "0"))
                    for a in tgt_root.findall(f"{{{_W}}}abstractNum")]
    existing_num = [int(n.get(f"{{{_W}}}numId", "0"))
                    for n in tgt_root.findall(f"{{{_W}}}num")]
    abs_offset = max(existing_abs, default=0) + 1
    num_offset = max(existing_num, default=0) + 1

    abs_map = {}
    for i, abst in enumerate(src_root.findall(f"{{{_W}}}abstractNum")):
        old_id = abst.get(f"{{{_W}}}abstractNumId")
        new_id = str(abs_offset + i)
        abs_map[old_id] = new_id
        cloned = deepcopy(abst)
        cloned.set(f"{{{_W}}}abstractNumId", new_id)
        first_num = tgt_root.find(f"{{{_W}}}num")
        if first_num is not None:
            tgt_root.insert(list(tgt_root).index(first_num), cloned)
        else:
            tgt_root.append(cloned)

    num_map = {}
    for i, num in enumerate(src_root.findall(f"{{{_W}}}num")):
        old_nid = num.get(f"{{{_W}}}numId")
        new_nid = str(num_offset + i)
        num_map[old_nid] = new_nid
        cloned = deepcopy(num)
        cloned.set(f"{{{_W}}}numId", new_nid)
        ref = cloned.find(f"{{{_W}}}abstractNumId")
        if ref is not None:
            ref.set(f"{{{_W}}}val", abs_map.get(ref.get(f"{{{_W}}}val"), ref.get(f"{{{_W}}}val")))
        tgt_root.append(cloned)

    return num_map


def _resize_inline_images(cloned_elem, max_w_in: float = 5.5, max_h_in: float = 6.5):
    """
    Rescale any inline picture inside a cloned paragraph so it fits within
    max_w_in x max_h_in, preserving aspect ratio. Only shrinks - never
    enlarges. Updates both wp:extent (drawing frame) and the matching
    a:ext inside the picture's xfrm so Word renders them consistently.
    """
    _WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    _A  = "http://schemas.openxmlformats.org/drawingml/2006/main"
    for inline in (cloned_elem.findall(f".//{{{_WP}}}inline")
                   + cloned_elem.findall(f".//{{{_WP}}}anchor")):
        ext = inline.find(f"{{{_WP}}}extent")
        if ext is None:
            continue
        try:
            cx = int(ext.get("cx")); cy = int(ext.get("cy"))
        except (TypeError, ValueError):
            continue
        if cx <= 0 or cy <= 0:
            continue
        w_in, h_in = cx / 914400, cy / 914400
        if w_in <= max_w_in and h_in <= max_h_in:
            continue  # already fits — leave untouched
        aspect    = cx / cy
        new_w_in  = min(max_w_in, max_h_in * aspect)
        new_h_in  = new_w_in / aspect
        new_cx, new_cy = int(new_w_in * 914400), int(new_h_in * 914400)
        ext.set("cx", str(new_cx)); ext.set("cy", str(new_cy))
        for xfrm_ext in inline.findall(f".//{{{_A}}}ext"):
            xfrm_ext.set("cx", str(new_cx)); xfrm_ext.set("cy", str(new_cy))


def _clone_para(src_elem, num_map: dict = None, full_para_text: str = None,
                 max_img_w_in: float = 5.5, max_img_h_in: float = 6.5,
                 strip_drawings: bool = False):
    """
    Deep-clone a paragraph applying MOP house style.
    num_map: remaps source numIds to imported IDs so numbering format is
             preserved exactly from the source document.
    strip_drawings: if True, remove all w:drawing elements from the clone.
                    Used for MOP paragraphs whose images are handled separately
                    via the MediaItem queue to prevent image duplication.
    """
    cloned = deepcopy(src_elem)

    # ── Strip inline drawings from MOP-sourced paragraphs ─────────────────────
    # When a MOP paragraph that contains an inline image is cloned via deepcopy,
    # the <w:drawing> XML (with its rId relationship) is carried verbatim into
    # the output document body.  At the same time, extract_media_from_activity_mop
    # has already extracted the same image blob into the MediaItem queue, which
    # later re-inserts it via _make_image_xml / _make_image_para.
    # The result is the image appearing TWICE — once from the cloned paragraph
    # XML and once from the media queue re-insertion.
    # Fix: when strip_drawings=True, remove <w:drawing> runs entirely so only
    # the media-queue path produces the image in the document.
    if strip_drawings:
        _WD = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        _WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        # Remove runs that contain only a drawing (no text)
        for run in list(cloned.findall(f".//{{{_WD}}}r")):
            if run.find(f"{{{_WD}}}drawing") is not None:
                run_texts = "".join(t.text or "" for t in run.findall(f"{{{_WD}}}t")).strip()
                if not run_texts:
                    parent = run.getparent()
                    if parent is not None:
                        parent.remove(run)
        # Also strip any stray <w:drawing> that sits directly inside a run
        for drawing in list(cloned.findall(f".//{{{_WD}}}drawing")):
            parent = drawing.getparent()
            if parent is not None:
                parent.remove(drawing)
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    pStyle_el  = cloned.find(".//" + qn("w:pStyle"))
    style_val  = pStyle_el.get(qn("w:val"), "") if pStyle_el is not None else ""
    is_heading = "Heading" in style_val

    if full_para_text is None:
        full_para_text = "".join(
            t.text or "" for t in cloned.findall(f".//{{{_W}}}t")
        ).strip()

    if num_map and not is_heading:
        numPr = cloned.find(".//" + qn("w:numPr"))
        if numPr is not None:
            nid_el = numPr.find(qn("w:numId"))
            if nid_el is not None:
                old_nid = nid_el.get(qn("w:val"), "0")
                nid_el.set(qn("w:val"), num_map.get(old_nid, old_nid))

    # ── Strip numPr from ALL Heading paragraphs ───────────────────────────────
    # Heading 2 in SOP (Phase 1, Phase 2 etc.) inherits numId from style
    # definition causing 1.1/1.2 numbering and left-shift in output.
    # Manual version has no numPr on headings — strip it to match.
    if is_heading:
        pPr_h = cloned.find(f"{{{_W}}}pPr")
        if pPr_h is not None:
            numPr_h = pPr_h.find(f"{{{_W}}}numPr")
            if numPr_h is not None:
                pPr_h.remove(numPr_h)

    # ── Ensure list paragraphs have explicit indent ───────────────────────────
    # When a numbered/bulleted paragraph is cloned into the template, its indent
    # may be inherited from "List Paragraph" style in the source doc but that
    # style may not be defined the same way in the template — causing left=0.
    # Fix: always set explicit ind on numPr paragraphs.
    if not is_heading:
        pPr_ind = cloned.find(f"{{{_W}}}pPr")
        if pPr_ind is not None:
            numPr_check = pPr_ind.find(f"{{{_W}}}numPr")
            if numPr_check is not None:
                ilvl_el = numPr_check.find(f"{{{_W}}}ilvl")
                ilvl_val = int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0
                ind_el = pPr_ind.find(f"{{{_W}}}ind")
                if ind_el is None:
                    ind_el = OxmlElement("w:ind")
                    pPr_ind.append(ind_el)
                # Set indent based on level: ilvl=0 → 360twips, ilvl=1 → 720twips
                base_left    = 360 + ilvl_val * 360
                base_hanging = 360
                ind_el.set(qn("w:left"),    str(base_left))
                ind_el.set(qn("w:hanging"), str(base_hanging))

    # ── Strip pBdr (paragraph border = separator line) from ALL paragraphs ──
    pPr_cl = cloned.find(f"{{{_W}}}pPr")
    if pPr_cl is not None:
        pBdr_el = pPr_cl.find(f"{{{_W}}}pBdr")
        if pBdr_el is not None:
            pPr_cl.remove(pBdr_el)

    # ── Strip footnote & endnote reference runs ───────────────────────────────
    # footnotes.xml / endnotes.xml from the source doc are never transferred to
    # the output template, so any w:footnoteReference or w:endnoteReference tags
    # that survive the deepcopy cause Word "unreadable content / Show Repairs"
    # errors on open.  These runs contain ZERO visible text (only a superscript
    # number marker) so removing them loses no content whatsoever.
    # Safe for normal docs too — findall returns [] when none exist → no-op.
    for _fn_run in cloned.findall(f".//{{{_W}}}r"):
        if (_fn_run.find(f"{{{_W}}}footnoteReference") is not None or
                _fn_run.find(f"{{{_W}}}endnoteReference") is not None):
            _fn_run.getparent().remove(_fn_run)

    for run in cloned.findall(f".//{{{_W}}}r"):
        rpr = run.find(f"{{{_W}}}rPr")
        if rpr is None:
            rpr = OxmlElement("w:rPr"); run.insert(0, rpr)
        run_text = "".join(t.text or "" for t in run.findall(f"{{{_W}}}t")).strip()
        if not is_heading:
            _normalise_run(rpr)
            b_el = rpr.find(f"{{{_W}}}b")
            if b_el is not None:
                is_label = (run_text and run_text == full_para_text
                            and len(run_text) < 80)
                if not is_label:
                    rpr.remove(b_el)
            for bcs in rpr.findall(f"{{{_W}}}bCs"): rpr.remove(bcs)

    _resize_inline_images(cloned, max_w_in=max_img_w_in, max_h_in=max_img_h_in)

    return cloned


def _force_table_full_width(tbl_elem):
    """
    Set a table to occupy 100% of the destination page's text width,
    preserving its existing column-ratio proportions exactly.
    Uses tblLayout="fixed" (not "autofit") because "autofit" tells
    Word/LibreOffice to size columns off cell *content* instead of the
    declared percentage — which left short-content tables (e.g. a 2-column
    table with just "FAN"/"Alarm Name") rendering narrower than the page
    even with tblW set to 100%. "fixed" honours the percentage width
    literally and distributes it across columns using their original
    relative ratios.
    """
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tblPr = tbl_elem.find(f"{{{_W}}}tblPr")
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    for tag in ("tblW", "tblLayout"):
        old = tblPr.find(f"{{{_W}}}{tag}")
        if old is not None:
            tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)


def _clone_table(src_tbl, num_map: dict = None):
    """Deep-clone a table normalising cell run formatting to house style."""
    cloned = deepcopy(src_tbl)
    _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for run in cloned.findall(f".//{{{_W}}}r"):
        rpr = run.find(f"{{{_W}}}rPr")
        if rpr is None:
            rpr = OxmlElement("w:rPr"); run.insert(0, rpr)
        _normalise_run(rpr)
        for b  in rpr.findall(f"{{{_W}}}b"):   rpr.remove(b)
        for bc in rpr.findall(f"{{{_W}}}bCs"): rpr.remove(bc)

    # ── Make table fill 100% of the destination page's text width ──────────
    # Source tables (Solution Doc / Activity MOP) carry an absolute width
    # (dxa) sized for whichever page they were authored on. Cloning that
    # verbatim into a template with a different page size/orientation makes
    # the table look too narrow (or overflow). See _force_table_full_width.
    _force_table_full_width(cloned)

    return cloned






# ─────────────────────────────────────────────────────────────────
# MAIN BUILD FUNCTION
# ─────────────────────────────────────────────────────────────────
def _embed_attachment_into_docx(doc_bytes: bytes, media_item, attach_idx: int):
    try:
        fname    = media_item.display_name
        ext      = media_item.ext.lower()
        CT_MAP   = {
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "xls":  "application/vnd.ms-excel",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "doc":  "application/msword",
            "txt":  "text/plain", "log": "text/plain", "csv": "text/csv",
            "pdf":  "application/pdf", "zip": "application/zip",
            "bin":  "application/octet-stream",
        }
        ct       = CT_MAP.get(ext, "application/octet-stream")
        rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject"
        new_rId  = f"rIdEmb{attach_idx:04d}"
        safe_name = re.sub(r'[^\w.\-]', '_', fname)
        part_path = f"word/embeddings/{safe_name}"
        rel_tgt   = f"embeddings/{safe_name}"
        in_buf, out_buf = io.BytesIO(doc_bytes), io.BytesIO()
        with zipfile.ZipFile(in_buf, "r") as zin, \
             zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                info = zipfile.ZipInfo(item.filename)
                info.compress_type = zipfile.ZIP_DEFLATED
                if item.filename == "word/_rels/document.xml.rels":
                    rels = data.decode("utf-8")
                    new_rel = f'<Relationship Id="{new_rId}" Type="{rel_type}" Target="{rel_tgt}"/>'
                    rels = rels.replace("</Relationships>", new_rel + "</Relationships>")
                    zout.writestr(info, rels.encode("utf-8"))
                elif item.filename == "[Content_Types].xml":
                    cts = data.decode("utf-8")
                    if part_path not in cts:
                        cts = cts.replace("</Types>",
                            f'<Override PartName="/{part_path}" ContentType="{ct}"/></Types>')
                    zout.writestr(info, cts.encode("utf-8"))
                else:
                    zout.writestr(info, data)
            zout.writestr(part_path, media_item.blob)
        out_buf.seek(0)
        return out_buf.read(), True
    except Exception:
        return doc_bytes, False



def build_mop(
    template_bytes:  bytes,
    activity_name:   str,
    sections:        dict,
    today_str:       str,
    media_items:     list,
    mop_sections:    dict = None,
) -> tuple:
    doc  = Document(io.BytesIO(template_bytes))
    body = doc.element.body
    _update_header_date(doc, today_str)

    # ── Force portrait orientation ──────────────────────────────────────────
    # Generated MOPs should always be portrait, regardless of whichever
    # template (landscape or portrait) was loaded. Only page_width/height +
    # the orientation flag are swapped — margins are left exactly as
    # authored in the template, since they may be deliberately tuned
    # (e.g. a tall top margin for a letterhead banner).
    _orientation_swapped = False
    for _sec in doc.sections:
        if _sec.orientation == WD_ORIENT.LANDSCAPE or _sec.page_width > _sec.page_height:
            _w, _h = _sec.page_width, _sec.page_height
            _sec.page_width, _sec.page_height = _h, _w
            _sec.orientation = WD_ORIENT.PORTRAIT
            _orientation_swapped = True

    # A header/footer/cover-page table already in the template (not one we
    # clone in later) carries an absolute width sized for the page it was
    # authored on. If we just narrowed the page (landscape → portrait),
    # those static tables would now overflow the new, narrower width —
    # so re-fit them too, the same way cloned tables get re-fit.
    if _orientation_swapped:
        for _sec in doc.sections:
            for _tbl in list(_sec.header.tables) + list(_sec.footer.tables):
                _force_table_full_width(_tbl._tbl)
        for _tbl in doc.tables:
            _force_table_full_width(_tbl._tbl)

    # ── Derive real page-fit limits for inserted images from the actual
    # template page size/margins/orientation (not a generic hardcoded guess).
    # A landscape template with a tall top margin (letterhead/branding) can
    # have very little usable height — sizing purely off a portrait
    # assumption would still overflow such templates.
    try:
        _sect    = doc.sections[0]
        _usable_w_in = (_sect.page_width.twips - _sect.left_margin.twips
                        - _sect.right_margin.twips) / 1440
        _usable_h_in = (_sect.page_height.twips - _sect.top_margin.twips
                         - _sect.bottom_margin.twips) / 1440
        IMG_MAX_W = max(min(5.5, _usable_w_in - 0.3), 2.0)
        # Leave headroom for heading text above + caption below the image.
        IMG_MAX_H = max(_usable_h_in * 0.75, 2.0)
    except Exception:
        IMG_MAX_W, IMG_MAX_H = 5.5, 6.5

    # ── Strip numPr from Heading2–9 style definitions in template ────────────
    # Template has numId=1, ilvl=1..8 in Heading2–9 styles which causes Word
    # to render 1.1, 1.2... numbering on all subheadings (Phase 1, Phase 2).
    # Fix: remove numPr from these style definitions at runtime so headings
    # always render without auto-numbering, aligned at left=0.
    _W_S = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        styles_part = doc.part.styles
        for style in styles_part._element.findall(f".//{{{_W_S}}}style"):
            sid = style.get(f"{{{_W_S}}}styleId", "")
            if sid in ("Heading2","Heading3","Heading4","Heading5",
                       "Heading6","Heading7","Heading8","Heading9"):
                pPr = style.find(f"{{{_W_S}}}pPr")
                if pPr is not None:
                    numPr = pPr.find(f"{{{_W_S}}}numPr")
                    if numPr is not None:
                        pPr.remove(numPr)
    except Exception:
        pass

    _update_revision_table(doc, activity_name, today_str)

    # ── Remove borders ONLY from the cover table (contains MOP title) ──────────
    # The cover/title table has visible borders that render as dark lines.
    # The revision history table MUST keep its borders and blue header.
    # Strategy: only strip borders from a table whose first cell contains
    # "METHOD OF PROCEDURE" or is blank/title-only (i.e. the cover table).
    _W_TBL = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for tbl in doc.tables:
        # Check if this is the cover table
        first_cell_txt = ""
        rows = tbl._tbl.findall(f".//{{{_W_TBL}}}tr")
        if rows:
            cells = rows[0].findall(f".//{{{_W_TBL}}}tc")
            if cells:
                first_cell_txt = "".join(
                    t.text or "" for t in cells[0].findall(f".//{{{_W_TBL}}}t")
                ).strip().upper()

        # Only remove borders if it looks like the cover/title table
        # (contains "METHOD" or "PROCEDURE" or "TITLE", or is a 1-cell table)
        is_cover = (
            "METHOD" in first_cell_txt or
            "PROCEDURE" in first_cell_txt or
            "TITLE" in first_cell_txt or
            (len(rows) == 1 and sum(1 for _ in rows[0].findall(f".//{{{_W_TBL}}}tc")) <= 1)
        )
        if not is_cover:
            continue   # keep revision table and all other tables intact

        tblPr = tbl._tbl.find(f"{{{_W_TBL}}}tblPr")
        if tblPr is not None:
            tblBdr = tblPr.find(f"{{{_W_TBL}}}tblBorders")
            if tblBdr is not None:
                tblPr.remove(tblBdr)
        for cell in tbl._tbl.findall(f".//{{{_W_TBL}}}tc"):
            tcPr = cell.find(f"{{{_W_TBL}}}tcPr")
            if tcPr is not None:
                tcBdr = tcPr.find(f"{{{_W_TBL}}}tcBorders")
                if tcBdr is not None:
                    tcPr.remove(tcBdr)

    # ── Activity name subtitle under title ────────────────────────────────────
    for child in list(body):
        if child.tag.split("}")[-1] != "p":
            continue
        se = child.find(".//" + qn("w:pStyle"))
        if se is not None and se.get(qn("w:val")) == "Title":
            # Force center alignment on the Title paragraph itself
            pPr_title = child.find(qn("w:pPr"))
            if pPr_title is None:
                pPr_title = OxmlElement("w:pPr"); child.insert(0, pPr_title)
            jc_title = pPr_title.find(qn("w:jc"))
            if jc_title is None:
                jc_title = OxmlElement("w:jc"); pPr_title.append(jc_title)
            jc_title.set(qn("w:val"), "center")

            sub_e = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            jc  = OxmlElement("w:jc"); jc.set(qn("w:val"), "center")
            sp  = OxmlElement("w:spacing"); sp.set(qn("w:before"), "80"); sp.set(qn("w:after"), "80")
            pPr.append(jc); pPr.append(sp)
            sub_e.append(pPr)
            r   = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            fn  = OxmlElement("w:rFonts"); fn.set(qn("w:ascii"), _BODY_FONT); fn.set(qn("w:hAnsi"), _BODY_FONT)
            sz  = OxmlElement("w:sz");  sz.set(qn("w:val"), "28")
            szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "28")
            it  = OxmlElement("w:i"); itCs = OxmlElement("w:iCs")
            col = OxmlElement("w:color"); col.set(qn("w:val"), "000000")
            rpr.append(fn); rpr.append(sz); rpr.append(szCs)
            rpr.append(it); rpr.append(itCs); rpr.append(col)
            r.append(rpr)
            t   = OxmlElement("w:t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = activity_name
            r.append(t); sub_e.append(r)
            _insert_after(child, sub_e)
            break

    # Import numbering definitions from both source docs into template.
    sol_num_map = _import_numbering_defs(
        build_mop._sol_bytes, doc) if getattr(build_mop, '_sol_bytes', None) else {}
    mop_num_map = _import_numbering_defs(
        build_mop._mop_bytes, doc) if getattr(build_mop, '_mop_bytes', None) else {}

    # Remove template separator paragraphs — empty <w:p> with only a pBdr border
    # These appear as dark horizontal lines between sections in the output document
    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for child in list(body):
        if child.tag.split("}")[-1] != "p":
            continue
        text = "".join(t.text or "" for t in child.findall(f".//{{{_W_NS}}}t")).strip()
        if not text:
            pPr = child.find(f"{{{_W_NS}}}pPr")
            if pPr is not None:
                pBdr = pPr.find(f"{{{_W_NS}}}pBdr")
                if pBdr is not None:
                    body.remove(child)

    # Collect ordered Heading1 elements — skip blank ones (separator artifacts)
    ordered_sections = []
    for child in list(body):
        if child.tag.split("}")[-1] != "p":
            continue
        se = child.find(".//" + qn("w:pStyle"))
        if se is None or se.get(qn("w:val"), "") != "Heading1":
            continue
        text = "".join(r.text or "" for r in child.findall(".//" + qn("w:t"))).strip()
        if not text:
            body.remove(child)   # blank Heading1 = template separator line
            continue
        key = normalize_heading(text)
        if key:
            _apply_heading_color(child)
            ordered_sections.append((child, key))

    media_queue    = list(media_items)
    media_idx      = 0
    failed_media   = []
    injected_count = 0
    pending_att    = []
    att_counter    = [0]

    for h_elem, sec_key in ordered_sections:

        # ── Clear template boilerplate under heading ──────────────────────────
        # Stop ONLY at the next Heading1 — Heading2/3 are template boilerplate
        # that must also be cleared (e.g. "Phase 1/2" stub under SOP heading).
        to_remove, found = [], False
        for child in list(body):
            if child is h_elem:
                found = True; continue
            if not found: continue
            ctag = child.tag.split("}")[-1]
            if ctag == "sectPr": break
            if ctag == "p":
                se = child.find(".//" + qn("w:pStyle"))
                sv = se.get(qn("w:val"), "") if se is not None else ""
                # Stop only at Heading1 — never at Heading2/3 (those are boilerplate)
                if sv in ("Heading1", "1") or sv.startswith("Heading 1"):
                    break
                to_remove.append(child)
            elif ctag == "tbl":
                to_remove.append(child)
        for e in to_remove:
            body.remove(e)

        if sec_key == "objective":
            pPr = h_elem.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr"); h_elem.insert(0, pPr)
            pb = OxmlElement("w:pageBreakBefore"); pb.set(qn("w:val"), "1")
            pPr.append(pb)

        content_elems = sections.get(sec_key, [])
        mop_content   = (mop_sections or {}).get(sec_key, [])

        if not content_elems and not mop_content:
            _insert_after(h_elem, OxmlElement("w:p"))
            continue

        cursor = h_elem

        # ── Universal element inserter (num_map passed via closure) ───────────
        def _insert_elem(elem, cur, _snm=sol_num_map, _mnm=mop_num_map,
                         _is_mop=False):
            ctag = elem.tag.split("}")[-1]
            nm   = _mnm if _is_mop else _snm
            if ctag == "tbl":
                tbl_clone = _clone_table(elem, nm)
                parent    = cur.getparent()
                idx       = list(parent).index(cur)
                parent.insert(idx + 1, tbl_clone)
                return list(parent)[idx + 1]
            else:
                full_txt = "".join(
                    t.text or "" for t in elem.findall(".//" + qn("w:t"))
                ).strip()
                # strip_drawings=True for MOP elements: images from MOP are
                # re-inserted via the MediaItem queue to avoid duplication.
                cl = _clone_para(elem, num_map=nm, full_para_text=full_txt,
                                  max_img_w_in=IMG_MAX_W, max_img_h_in=IMG_MAX_H,
                                  strip_drawings=_is_mop and bool(media_items))
                _insert_after(cur, cl)
                return cl

        if sec_key == "sop":
            seen_texts = set()

            # ── Helper to inject one MediaItem at the current cursor ──────────
            def _inject_media_item(m, cur):
                nonlocal injected_count
                if m.kind == "image":
                    try:
                        img_xml = _make_image_xml(doc, m.blob,
                                                   max_w=IMG_MAX_W, max_h=IMG_MAX_H)
                        _insert_after(cur, img_xml); cur = img_xml
                        cap = _make_caption_xml()
                        _insert_after(cur, cap); cur = cap
                        m.injected = True; injected_count += 1
                    except Exception as ex:
                        m.inject_error = str(ex)
                        ctx = m.context_text
                        img_desc = (f"Image: '{ctx[:60]}'" if ctx and len(ctx) > 3
                                    else f"Image #{m.position_index+1}")
                        failed_media.append(img_desc)
                        _insert_after(cur, _make_notice_xml(img_desc))
                        cur = list(body)[-1]
                else:
                    att_counter[0] += 1
                    att_xml = _make_xml_para(
                        doc,
                        f"\U0001f4ce  ATTACHED FILE [{m.ext.upper()}]: "
                        f"{m.display_name}  \u2014 embedded in document",
                        bold=False, color_rgb="000000", size_pt=11)
                    _insert_after(cur, att_xml); cur = att_xml
                    pending_att.append((m, att_counter[0]))
                    m.injected = True; injected_count += 1
                return cur

            # ── NS constants for drawing detection ────────────────────────────
            _WD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            # ── Step A: Insert MOP content, injecting images at their original
            #    position (right after the para where the drawing was stripped).
            #    Each MOP paragraph that HAD a <w:drawing> in the source is
            #    matched 1-to-1 with the next MediaItem from the queue.
            if mop_sections:
                for elem in mop_sections.get("sop", []):
                    tag = elem.tag.split("}")[-1]
                    if tag == "tbl":
                        cursor = _insert_elem(elem, cursor, _is_mop=True)
                        continue
                    txt = "".join(t.text or "" for t in elem.findall(".//" + qn("w:t"))).strip()
                    if txt: seen_texts.add(txt.lower())

                    # Check if this source element contained an inline drawing
                    had_drawing = bool(elem.findall(
                        f".//{{{_WD_NS}}}drawing"))

                    cursor = _insert_elem(elem, cursor, _is_mop=True)

                    # If it had a drawing, inject the next media item HERE
                    # (the drawing XML was stripped from the clone by _clone_para,
                    #  so the image must come from the media queue instead)
                    if had_drawing and media_idx < len(media_queue):
                        media_item = media_queue[media_idx]; media_idx += 1
                        cursor = _inject_media_item(media_item, cursor)

            # ── Step B: Insert solution-doc SOP content (deduped) ────────────
            for p_elem in content_elems:
                tag = p_elem.tag.split("}")[-1]
                if tag == "tbl":
                    cursor = _insert_elem(p_elem, cursor); continue
                txt = "".join(t.text or "" for t in p_elem.findall(".//" + qn("w:t"))).strip()

                # Check for [IMAGE] placeholder in solution doc content
                is_ph = bool(IMAGE_PLACEHOLDER_RE.search(txt))

                if txt and txt.lower() in seen_texts: continue
                if txt: seen_texts.add(txt.lower())
                cursor = _insert_elem(p_elem, cursor)

                # Inject media at solution-doc placeholders (when no MOP uploaded)
                if is_ph and not mop_sections and media_idx < len(media_queue):
                    media_item = media_queue[media_idx]; media_idx += 1
                    cursor = _inject_media_item(media_item, cursor)

            # ── Step C: Remaining unmatched media → append at SOP end ─────────
            while media_idx < len(media_queue):
                m = media_queue[media_idx]; media_idx += 1
                cursor = _inject_media_item(m, cursor)


        elif sec_key == "acceptance_criteria":
            # ── UAT: copy as-is from solution doc — same rule as all other sections ──
            # Solution doc already has correct format (List Paragraph with w:br line breaks
            # containing Title + Condition/Measurement/Pass/Fail within one paragraph).
            # No custom logic needed — just clone and insert preserving structure.
            seen_texts = set()

            for elem in mop_content:
                txt = "".join(t.text or "" for t in elem.findall(".//" + qn("w:t"))).strip()
                if txt: seen_texts.add(txt.lower())
                cursor = _insert_elem(elem, cursor, _is_mop=True)

            for p_elem in content_elems:
                tag = p_elem.tag.split("}")[-1]
                if tag == "tbl":
                    cursor = _insert_elem(p_elem, cursor); continue
                txt = "".join(t.text or "" for t in p_elem.findall(".//" + qn("w:t"))).strip()
                if txt and txt.lower() in seen_texts: continue
                if txt: seen_texts.add(txt.lower())
                cursor = _insert_elem(p_elem, cursor)

        else:
            # ── Generic sections: MOP content first, then solution doc ────────
            seen_texts = set()

            for elem in mop_content:
                txt = "".join(t.text or "" for t in elem.findall(".//" + qn("w:t"))).strip()
                if txt: seen_texts.add(txt.lower())
                cursor = _insert_elem(elem, cursor, _is_mop=True)

            for p_elem in content_elems:
                tag = p_elem.tag.split("}")[-1]
                if tag == "tbl":
                    cursor = _insert_elem(p_elem, cursor); continue
                txt = "".join(t.text or "" for t in p_elem.findall(".//" + qn("w:t"))).strip()
                if txt and txt.lower() in seen_texts: continue
                if txt: seen_texts.add(txt.lower())
                cursor = _insert_elem(p_elem, cursor)

    # ── Final cleanup: remove ALL pBdr borders everywhere ──────────────────────
    # This strips separator lines from ALL paragraphs including Title style.
    # User rule: no lines anywhere in the output document.
    _W_FINAL = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _body_final = doc.element.body
    for _child in list(_body_final):
        if _child.tag.split("}")[-1] != "p":
            continue
        _pPr = _child.find("{%s}pPr" % _W_FINAL)
        if _pPr is not None:
            _pBdr = _pPr.find("{%s}pBdr" % _W_FINAL)
            if _pBdr is not None:
                _pPr.remove(_pBdr)

    # Save base document
    buf = io.BytesIO()
    doc.save(buf)
    doc_bytes = buf.getvalue()

    # Post-process: embed each attachment into the docx ZIP
    for media_item, att_idx in pending_att:
        new_bytes, ok = _embed_attachment_into_docx(doc_bytes, media_item, att_idx)
        if ok:
            doc_bytes = new_bytes
        else:
            desc = (f"Attachment '{media_item.display_name}' (.{media_item.ext}) "
                    f"could not be embedded — please attach manually")
            failed_media.append(desc)

    return doc_bytes, failed_media, injected_count


# ─────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style="text-align:center; padding:1.2rem 0 0.8rem;">
      <span style="font-family:'Lato',sans-serif; font-weight:900; font-size:2rem; letter-spacing:6px; color:#0082C8; display:block;">ERICSSON</span>
      <span style="font-size:0.6rem; letter-spacing:2px; color:rgba(255,255,255,0.3); text-transform:uppercase; display:block; margin-top:2px;">Technology For Good</span>
    </div>
    <hr/>
    """, unsafe_allow_html=True)

    st.markdown("### 📋 Smart MOP Generator")
    st.markdown("""
    <div class="sidebar-info">
      Unified solution for all MOP generation scenarios.<br><br>
      <strong>Supported inputs:</strong><br>
      · Activity MOP only<br>
      · Transcript only<br>
      · Transcript + Logs<br>
      · Transcript + Activity MOP<br>
      · Activity MOP + Images<br>
      · Logs only<br>
      · Any mixed combination
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("**Workflow:**")
    st.markdown("""
    <div style="font-size:0.78rem; color:#90b8e0; line-height:1.9;">
    1️⃣ &nbsp;Use <strong>Master Prompt</strong> in Copilot<br>
    &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;with your input files<br>
    2️⃣ &nbsp;Save Copilot output as<br>
    &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;<strong>Solution Document (.docx)</strong><br>
    3️⃣ &nbsp;Upload Solution Document here<br>
    4️⃣ &nbsp;Optionally upload Activity MOP<br>
    &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;for image/attachment injection<br>
    5️⃣ &nbsp;Optionally upload extra files<br>
    &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;(logs, screenshots, tables…)<br>
    6️⃣ &nbsp;Click <strong>Generate Smart MOP</strong><br>
    7️⃣ &nbsp;Download final <strong>.docx</strong>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("""
    <div class="sidebar-info">
      🔒 <strong>Zero Data Retention</strong><br>
      All processing strictly in-memory.<br>
      No files written to disk.<br>
      No data logged or stored.<br>
      Session cleared on browser close.
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("""
    <div style="font-size:0.62rem; color:rgba(255,255,255,0.2); text-align:center; letter-spacing:.5px;">
      Smart MOP Generator v1.0 — Unified<br>
      Ericsson Internal Tool
    </div>
    """, unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────
# MAIN AREA
# ─────────────────────────────────────────────────────────────────
st.markdown("""
<div class="eri-topbar">
  <div>
    <div class="eri-logo-text">ERICSSON</div>
    <div class="eri-logo-sub">Telecom Automation Toolkit</div>
  </div>
  <div style="text-align:center;">
    <div class="eri-app-title">Smart MOP Generator</div>
    <div class="eri-app-sub">Unified · All Scenarios · Any Input Language → Professional English MOP · Audit-Ready · ZDR</div>
  </div>
  <div>
    <span class="eri-version">v1.0 UNIFIED</span>
  </div>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="priv-bar">
  🔒 <strong>ZERO DATA RETENTION:</strong> All processing is performed entirely in-memory.
  No uploaded files, generated documents, or any user data are written to disk, logged, or stored
  at any stage. This session and all associated data are permanently cleared when you close your browser.
</div>
""", unsafe_allow_html=True)

# ── Layout ──────────────────────────────────────────────────────
col_left, col_right = st.columns([1.1, 1], gap="large")

with col_left:

    # ── Step 1: Template ─────────────────────────────────────────
    st.markdown('<div class="eri-card"><div class="eri-card-title"><span class="step-badge">STEP 01</span> Select MOP Template</div>', unsafe_allow_html=True)

    templates         = discover_templates()
    selected_template = None
    template_bytes    = None

    if not templates:
        st.markdown('<div class="pill-warn">⚠ No template found. Place <strong>.docx</strong> file in <code>templates/</code> folder, then restart.</div>', unsafe_allow_html=True)
    else:
        names = [t.name for t in templates]
        sel   = st.selectbox("Template file", names, label_visibility="visible")
        selected_template = next(t for t in templates if t.name == sel)
        template_bytes    = load_template_bytes(selected_template)
        st.markdown(f'<div class="pill-ok">✔ &nbsp;<strong>{sel}</strong> &nbsp;ready</div>', unsafe_allow_html=True)

    st.markdown("""
    <div style="font-size:0.72rem; color:#5a7a9a; background:rgba(0,82,130,0.05);
         border:1px solid rgba(0,130,200,0.15); border-left:3px solid #0082C8;
         border-radius:0 6px 6px 0; padding:0.5rem 0.8rem; margin-top:0.5rem;">
      📁 &nbsp;<strong>To add/update a template</strong>, place the <code>.docx</code>
      file in the <code>templates/</code> folder and restart the app.
    </div>
    """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Step 2: Solution Document ────────────────────────────────
    st.markdown('<div class="eri-card"><div class="eri-card-title"><span class="step-badge">STEP 02</span> Upload Solution Document <span style="font-size:0.62rem;color:#cc4400;margin-left:6px;">REQUIRED</span></div>', unsafe_allow_html=True)
    st.markdown("""
    <div style="font-size:0.73rem;color:#5a7a9a;margin-bottom:0.5rem;">
      The Copilot-generated Solution Document (all 12 sections). This is the primary text source
      for your Smart MOP. Can come from <strong>any scenario</strong>: Activity MOP,
      transcript, logs, or any combination processed through the Master Prompt.
    </div>
    """, unsafe_allow_html=True)

    sol_file = st.file_uploader("Solution Document (.docx)", type=["docx"],
                                key="sol_up", label_visibility="visible")
    if sol_file:
        size_kb = sol_file.size / 1024
        st.markdown(
            f'<div class="pill-ok">✔ &nbsp;<strong>{sol_file.name}</strong>'
            f' &nbsp;·&nbsp; {size_kb:.1f} KB</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown('<div class="pill-warn">⏳ Waiting for solution document…</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Step 3: Activity MOP (OPTIONAL) ─────────────────────────
    st.markdown('<div class="eri-card"><div class="eri-card-title"><span class="step-badge">STEP 03</span> Upload Activity MOP &nbsp;<span class="optional-badge">OPTIONAL</span></div>', unsafe_allow_html=True)
    st.markdown("""
    <div style="font-size:0.73rem;color:#5a7a9a;margin-bottom:0.5rem;">
      The original Activity MOP document. <strong>Only media</strong> is extracted from this file
      — screenshots, flow diagrams, embedded attachments (Excel, logs, etc.) and Word comments.
      These are automatically injected into the SOP section of the output MOP at
      <code>[IMAGE/SCREENSHOT REQUIRED]</code> placeholder positions.
    </div>
    """, unsafe_allow_html=True)

    mop_file = st.file_uploader("Activity MOP (.docx) — optional", type=["docx"],
                                key="mop_up", label_visibility="visible")
    if mop_file:
        size_kb = mop_file.size / 1024
        st.markdown(
            f'<div class="pill-ok">✔ &nbsp;<strong>{mop_file.name}</strong>'
            f' &nbsp;·&nbsp; {size_kb:.1f} KB — images, attachments & comments will be extracted</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown('<div class="pill-info">ℹ No Activity MOP uploaded — output will contain text only with placeholder notices.</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Step 4: Generate ─────────────────────────────────────────
    st.markdown('<div class="eri-card"><div class="eri-card-title"><span class="step-badge">STEP 04</span> Generate Smart MOP</div>', unsafe_allow_html=True)

    can_go  = bool(sol_file and templates)
    gen_btn = st.button("⚡  Generate Smart MOP", disabled=not can_go)

    if not can_go:
        missing = []
        if not templates:
            missing.append("MOP template")
        if not sol_file:
            missing.append("solution document")
        if missing:
            st.markdown(f'<div class="pill-warn">⏳ Still needed: {" + ".join(missing)}</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)


with col_right:

    if gen_btn and can_go:
        has_mop   = bool(mop_file)
        steps = [
            "Loading MOP template",
            "Reading solution document",
            "Extracting activity name & date",
            "Parsing all 12 sections",
            "Reading Activity MOP" if has_mop else "No Activity MOP — text-only mode",
            "Extracting images & attachments" if has_mop else "Skipping media extraction",
            "Extracting Word comments" if has_mop else "No comments to extract",
            "Building Smart MOP document",
            "Injecting comments into SOP section",
            "Finalising & quality check",
        ]

        st.markdown('<div class="eri-card"><div class="eri-card-title">⚙ Processing</div>', unsafe_allow_html=True)
        st.markdown('<div class="prog-wrap">', unsafe_allow_html=True)
        phs = [st.empty() for _ in steps]
        for ph, s in zip(phs, steps):
            ph.markdown(f'<div class="ps wait"><div class="pd wait"></div>{s}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        try:
            activity_name   = ""
            sections        = {}
            today_str       = ""
            output_bytes    = b""
            media_items     = []
            comment_items   = []
            comments_xml    = ""
            comments_ext    = ""
            comments_ids    = ""
            comments_exs    = ""
            failed_media    = []
            injected_count  = 0
            c_injected      = []
            c_failed        = []
            extra_ref_names = []
            tmpl_b          = None
            mop_sections    = None

            for i, (ph, step) in enumerate(zip(phs, steps)):
                ph.markdown(f'<div class="ps doing"><div class="pd doing"></div>{step}…</div>', unsafe_allow_html=True)
                time.sleep(0.10)

                if i == 0:
                    tmpl_b = load_template_bytes(selected_template)

                elif i == 1:
                    sol_bytes = sol_file.read()
                    sol_doc   = Document(io.BytesIO(sol_bytes))

                elif i == 2:
                    activity_name = extract_activity_name(sol_doc)
                    today_str     = datetime.today().strftime("%d-%m-%Y")

                elif i == 3:
                    sections = extract_sections(sol_doc)

                elif i == 4:
                    if has_mop:
                        mop_bytes_data = mop_file.read()

                elif i == 5:
                    if has_mop:
                        media_items  = extract_media_from_activity_mop(mop_bytes_data)
                        mop_sections = extract_mop_sections(mop_bytes_data)

                elif i == 6:
                    if has_mop:
                        (comment_items, comments_xml,
                         comments_ext, comments_ids,
                         comments_exs) = extract_comments_from_mop(mop_bytes_data)

                elif i == 7:
                    # Pass source bytes so build_mop can import numbering defs
                    build_mop._sol_bytes = sol_bytes
                    build_mop._mop_bytes = mop_bytes_data if has_mop else None
                    output_bytes, failed_media, injected_count = build_mop(
                        tmpl_b, activity_name, sections, today_str,
                        media_items,
                        mop_sections if has_mop else None
                    )

                elif i == 9:
                    if comment_items and comments_xml:
                        output_bytes, c_injected, c_failed = _inject_comments_into_docx(
                            output_bytes, comment_items,
                            comments_xml, comments_ext,
                            comments_ids, comments_exs
                        )

                ph.markdown(f'<div class="ps done"><div class="pd done"></div>{step} ✓</div>', unsafe_allow_html=True)
                time.sleep(0.04)

            # ── Store in session_state ──
            st.session_state["output_bytes"]      = output_bytes
            st.session_state["activity_name"]     = activity_name
            st.session_state["today_str"]         = today_str
            st.session_state["sections"]          = sections
            st.session_state["filled"]            = sum(1 for k in SECTION_KEYS[:-1] if sections.get(k))
            st.session_state["images_n"]          = len([m for m in media_items if m.kind == "image"])
            st.session_state["total_n"]           = sum(len(v) for k, v in sections.items())
            st.session_state["failed_media"]      = failed_media
            st.session_state["injected_media"]    = injected_count
            st.session_state["comments_injected"] = len(c_injected)
            st.session_state["comments_failed"]   = c_failed

            st.markdown('</div>', unsafe_allow_html=True)

        except Exception as e:
            st.markdown('</div>', unsafe_allow_html=True)
            st.error(f"❌ Error during generation: {e}")
            import traceback
            st.code(traceback.format_exc())

    # ── Result panel ─────────────────────────────────────────────
    if st.session_state.get("output_bytes"):
        activity_name     = st.session_state["activity_name"]
        today_str         = st.session_state["today_str"]
        sections          = st.session_state["sections"]
        output_bytes      = st.session_state["output_bytes"]
        filled            = st.session_state["filled"]
        images_n          = st.session_state["images_n"]
        total_n           = st.session_state["total_n"]
        failed_media      = st.session_state["failed_media"]
        injected_media    = st.session_state["injected_media"]
        comments_injected = st.session_state.get("comments_injected", 0)
        comments_failed   = st.session_state.get("comments_failed", [])

        # ── Success card ─────────────────────────────────────────
        st.markdown(f"""
        <div class="success-card">
          <div class="success-icon">✅</div>
          <div class="success-title">Smart MOP Generated Successfully</div>
          <div class="success-sub">
            <strong class="success-name">{activity_name}</strong>
            &nbsp;·&nbsp; {today_str}
          </div>
        </div>""", unsafe_allow_html=True)

        safe_name = re.sub(r'[^\w\s\-]', '', activity_name).strip().replace(' ', '_')[:80]
        _dl_key   = f"dl_{abs(hash(safe_name + today_str)) % 10_000_000}"
        st.download_button(
            label="📥  Download Smart MOP (.docx)",
            data=io.BytesIO(output_bytes),
            file_name=f"{safe_name}_SmartMOP.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=_dl_key,
            use_container_width=True,
        )

        # ── Failed media warning ──────────────────────────────────
        if failed_media:
            st.markdown("""
            <div class="warn-bar">
              <strong>⚠ Some media items could not be automatically inserted.</strong>
              They are marked with a notice in the document at the exact position
              they should appear. Please insert them manually via Word's
              Insert → Pictures or Insert → Object.
            </div>
            """, unsafe_allow_html=True)
            st.markdown('<div class="media-fail-card">', unsafe_allow_html=True)
            st.markdown('<div class="media-fail-title">🖼 Media requiring manual insertion:</div>', unsafe_allow_html=True)
            for desc in failed_media:
                st.markdown(f'<div class="media-fail-item">• {desc}</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        elif images_n > 0:
            st.markdown(
                f'<div class="pill-ok" style="margin-top:0.6rem;">✔ &nbsp;All {injected_media} media items injected successfully</div>',
                unsafe_allow_html=True
            )

        if comments_injected > 0:
            st.markdown(
                f'<div class="pill-ok" style="margin-top:0.4rem;">'
                f'✔ &nbsp;{comments_injected} comment(s) injected (author + date preserved)</div>',
                unsafe_allow_html=True
            )
        if comments_failed:
            st.markdown("""
            <div class="warn-bar">
              <strong>⚠ Some comments could not be position-matched.</strong>
              Please add them manually via Word Review → New Comment.
            </div>""", unsafe_allow_html=True)



        # ── Metrics ───────────────────────────────────────────────
        st.markdown('<div class="eri-card" style="margin-top:0.8rem;"><div class="eri-card-title">📊 Generation Summary</div>', unsafe_allow_html=True)
        st.markdown(f"""
        <div class="metric-row">
          <div class="metric-box">
            <div class="metric-val">{filled}<span style="font-size:.85rem;color:#9aaab8;">/12</span></div>
            <div class="metric-sub">Sections Filled</div>
          </div>
          <div class="metric-box">
            <div class="metric-val">{injected_media}</div>
            <div class="metric-sub">Media Injected</div>
          </div>
          <div class="metric-box">
            <div class="metric-val">{comments_injected}</div>
            <div class="metric-sub">Comments Injected</div>
          </div>
          <div class="metric-box">
            <div class="metric-val">{len(failed_media) + len(comments_failed)}</div>
            <div class="metric-sub">Manual Fixes Needed</div>
          </div>
        </div>""", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        # ── Section fill status ───────────────────────────────────
        with st.expander("📋 Section fill status", expanded=False):
            for k in SECTION_KEYS[:-1]:
                label      = SECTION_LABELS.get(k, k)
                filled_flag = bool(sections.get(k))
                icon  = "✅" if filled_flag else "⚠️"
                color = "#006633" if filled_flag else "#cc5500"
                st.markdown(
                    f'<div style="font-size:0.77rem;color:{color};padding:3px 0;">'
                    f'{icon} &nbsp; {label}</div>',
                    unsafe_allow_html=True,
                )

    elif not st.session_state.get("output_bytes"):
        st.markdown("""
        <div class="eri-card" style="border:2px dashed #dde4ed; background:#fafbfc; min-height:380px;
          display:flex; flex-direction:column; align-items:center; justify-content:center; text-align:center;">
          <div style="font-size:3rem; margin-bottom:1rem;">📄</div>
          <div style="font-size:0.95rem; font-weight:700; color:#003366; margin-bottom:0.6rem;">
            Smart MOP output will appear here
          </div>
          <div style="font-size:0.76rem; color:#9aaab8; max-width:280px; line-height:1.7;">
            Upload your Copilot-generated Solution Document and click
            <strong>Generate Smart MOP</strong> to get started.<br><br>
            Optionally upload your Activity MOP to inject images,
            attachments, and comments automatically.
          </div>
        </div>
        """, unsafe_allow_html=True)

# ── Footer ────────────────────────────────────────────────────────
st.markdown("""
<div class="footer">
  Smart MOP Generator v1.0 Unified &nbsp;·&nbsp; Ericsson Internal Tool &nbsp;·&nbsp;
  🔒 Zero Data Retention &nbsp;·&nbsp; All processing in-memory only &nbsp;·&nbsp;
  All input languages → Professional English output
</div>
""", unsafe_allow_html=True)
