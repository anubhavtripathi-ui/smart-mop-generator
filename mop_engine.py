"""
Smart MOP Generator - Core Engine
"""
import os, re, zipfile
from datetime import date
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

HEADING_MAP = {
    "objective": ["objective","objectives","goal","goals","purpose","aim","overview","summary"],
    "activity":  ["activity description","activity","description","work description","task description","scope of work"],
    "type":      ["activity type","type of activity","work type","task type","change type","nature of activity"],
    "domain":    ["domain in scope","domain","scope","domains","network domain","technology domain"],
    "prereq":    ["pre-requisites","prerequisites","pre requisites","preconditions","prerequisite","preparation"],
    "inventory": ["inventory details","inventory","node details","equipment details","node inventory","asset details"],
    "connect":   ["node connectivity process","connectivity","node connectivity","connection process","access method"],
    "iam":       ["identity and access management","iam","access management","credentials","login details","authentication"],
    "trigger":   ["activity triggering method","trigger","triggering method","activity trigger","how to trigger"],
    "sop":       ["standard operating procedure","sop","procedure","operating procedure","detailed procedure"],
    "accept":    ["acceptance criteria","acceptance","uat scenarios","uat","validation criteria","success criteria"],
    "assume":    ["assumptions","assumption","assumed conditions","dependencies","notes","caveats"],
}

AUTO_FILL = {
    "objective": "The objective of this activity is to successfully execute the {activity} for {vendor}, ensuring minimal impact to network services and adherence to change management processes.",
    "activity":  "This document outlines the method of procedure for {activity} to be carried out by {vendor}. The activity involves planned steps to ensure seamless execution.",
    "type":      "Planned Maintenance Activity – {activity} ({vendor})",
    "domain":    "This activity covers the network domain relevant to {activity} as scoped and agreed with {vendor}.",
    "prereq":    "Prior to initiating {activity} with {vendor}, ensure all necessary approvals, maintenance windows, and access permissions are in place.",
    "inventory": "Node inventory details for {activity} with {vendor} to be confirmed prior to execution. Refer to the approved change request for specifics.",
    "connect":   "Connectivity to the relevant nodes for {activity} shall be established by {vendor} as per the approved access management process.",
    "iam":       "Access credentials for {activity} shall be provided by {vendor} in compliance with the Identity and Access Management policy.",
    "trigger":   "This activity will be triggered upon receipt of the approved change request and maintenance window confirmation for {activity} with {vendor}.",
    "accept":    "Acceptance criteria for {activity} with {vendor}: All configured services are operational, no alarms raised, sign-off obtained from stakeholders.",
    "assume":    "It is assumed all pre-requisites for {activity} with {vendor} are met, maintenance window is approved, and rollback procedure is available.",
}

JUNK_PATTERNS = [
    'this section is to provide', 'mention the', 'provide activity detailed',
    'provide fallback', 'please include a detailed description',
]

def normalize(text):
    return re.sub(r'\s+', ' ', text.strip().lower())

def match_heading(heading_text):
    norm = normalize(heading_text)
    best_key, best_len = None, 0
    for key, synonyms in HEADING_MAP.items():
        for syn in synonyms:
            if norm == syn:
                return key
            if norm.startswith(syn) and len(syn) > best_len:
                best_key, best_len = key, len(syn)
            if syn.startswith(norm) and len(norm) > 2 and len(norm) > best_len:
                best_key, best_len = key, len(norm)
    if not best_key:
        for key, synonyms in HEADING_MAP.items():
            for syn in synonyms:
                if len(syn) >= 10 and syn in norm and len(syn) > best_len:
                    best_key, best_len = key, len(syn)
    return best_key

def is_junk(text):
    t = text.strip().lower()
    if t.startswith('[') and t.endswith(']'):
        return True
    return any(j in t for j in JUNK_PATTERNS)

def has_real_content(elements, paragraphs):
    if any(p.strip() and not is_junk(p) for p in paragraphs):
        return True
    for elem in elements:
        xml = etree.tostring(elem).decode()
        if 'blip' in xml or 'OLEObject' in xml or 'oleObject' in xml or '<w:drawing>' in xml:
            return True
    return False

def extract_content_from_docx(input_path):
    doc = Document(input_path)
    content_map = {}
    current_key = None
    current_paragraphs = []
    current_elements = []

    for para in doc.paragraphs:
        style_name = para.style.name.lower()
        text = para.text.strip()
        xml_str = etree.tostring(para._element).decode()
        has_content = bool(text) or 'blip' in xml_str or 'OLEObject' in xml_str or '<w:drawing>' in xml_str

        if not has_content:
            if current_key:
                current_elements.append(deepcopy(para._element))
            continue

        matched_key = None
        if 'heading' in style_name:
            matched_key = match_heading(text)
        elif current_key is None and text:
            matched_key = match_heading(text)

        if matched_key:
            if current_key and has_real_content(current_elements, current_paragraphs):
                content_map[current_key] = {
                    'text': '\n'.join(p for p in current_paragraphs if not is_junk(p)),
                    'elements': current_elements
                }
            current_key = matched_key
            current_paragraphs = []
            current_elements = []
        else:
            if current_key:
                current_paragraphs.append(text)
                current_elements.append(deepcopy(para._element))

    if current_key and has_real_content(current_elements, current_paragraphs):
        content_map[current_key] = {
            'text': '\n'.join(p for p in current_paragraphs if not is_junk(p)),
            'elements': current_elements
        }
    return content_map, doc

def extract_content_from_txt(input_path):
    with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
        lines = f.readlines()
    content_map = {}
    current_key, current_lines = None, []
    for line in lines:
        text = line.strip()
        if not text:
            continue
        matched_key = match_heading(text.rstrip(':'))
        if matched_key:
            if current_key and current_lines:
                real = [l for l in current_lines if l.strip() and not is_junk(l)]
                if real:
                    content_map[current_key] = {'text': '\n'.join(real), 'elements': []}
            current_key = matched_key
            current_lines = []
        else:
            if current_key:
                current_lines.append(text)
    if current_key and current_lines:
        real = [l for l in current_lines if l.strip() and not is_junk(l)]
        if real:
            content_map[current_key] = {'text': '\n'.join(real), 'elements': []}
    return content_map

def copy_all_relationships(input_doc, output_doc, elements):
    """Copy ALL relationships (images PNG/EMF/JPG + OLE) from input to output."""
    rId_map = {}
    try:
        input_part  = input_doc.part
        output_part = output_doc.part
        # Get all rIds referenced in elements
        element_xml = ' '.join(etree.tostring(e).decode() for e in elements)
        flat_rids = set(re.findall(r'r:embed="(rId\d+)"|r:id="(rId\d+)"', element_xml))
        referenced = set()
        for pair in flat_rids:
            for rid in pair:
                if rid: referenced.add(rid)

        for rId in referenced:
            if rId not in input_part.rels:
                continue
            rel = input_part.rels[rId]
            try:
                target_part = rel.target_part
                new_rId = output_part.relate_to(target_part, rel.reltype)
                rId_map[rId] = new_rId
            except Exception:
                pass
    except Exception:
        pass
    return rId_map

def remap_rids_in_elements(elements, rId_map):
    if not rId_map:
        return
    r_embed = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    r_id    = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'
    for elem in elements:
        for node in elem.iter():
            for attr in [r_embed, r_id]:
                old = node.get(attr)
                if old and old in rId_map:
                    node.set(attr, rId_map[old])

def replace_placeholder_in_paragraph(para, placeholder, text):
    token = '{{' + placeholder + '}}'
    if token not in para.text:
        return False
    for run in para.runs:
        if token in run.text:
            run.text = run.text.replace(token, text)
            return True
    full_text = para.text.replace(token, text)
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = full_text
    return True

def replace_placeholders_in_doc(doc, replacements):
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if '{{' + key + '}}' in para.text:
                replace_placeholder_in_paragraph(para, key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in replacements.items():
                        if '{{' + key + '}}' in para.text:
                            replace_placeholder_in_paragraph(para, key, value)

def replace_content_for_key(output_doc, input_doc, key, content_data, activity_name, vendor_name):
    token = '{{' + key + '}}'
    placeholder_para = None
    for para in output_doc.paragraphs:
        if token in para.text:
            placeholder_para = para
            break
    if not placeholder_para:
        for table in output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if token in para.text:
                            placeholder_para = para
    if not placeholder_para:
        return

    elements = content_data.get('elements', [])
    text     = content_data.get('text', '').strip()

    if elements and input_doc:
        rId_map = copy_all_relationships(input_doc, output_doc, elements)
        if rId_map:
            remap_rids_in_elements(elements, rId_map)
        for run in placeholder_para.runs:
            run.text = run.text.replace(token, '')
        parent = placeholder_para._element.getparent()
        idx = list(parent).index(placeholder_para._element)
        for j, elem in enumerate(elements):
            parent.insert(idx + 1 + j, deepcopy(elem))
    elif text:
        replace_placeholder_in_paragraph(placeholder_para, key, text)
    else:
        replace_placeholder_in_paragraph(placeholder_para, key, '')

def _make_note_para(text, color='FF0000'):
    """Create a bold colored note paragraph."""
    note_para = OxmlElement('w:p')
    note_pPr  = OxmlElement('w:pPr')
    note_rPr_pPr = OxmlElement('w:rPr')
    note_pPr.append(note_rPr_pPr)
    note_para.append(note_pPr)
    note_r = OxmlElement('w:r')
    note_rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    col = OxmlElement('w:color')
    col.set(qn('w:val'), color)
    note_rPr.append(b)
    note_rPr.append(col)
    note_r.append(note_rPr)
    note_t = OxmlElement('w:t')
    note_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    note_t.text = text
    note_r.append(note_t)
    note_para.append(note_r)
    return note_para


def _copy_elements_with_rels(elements_to_copy, input_doc, output_doc, parent, insert_pos):
    """Copy body elements (p/tbl) remapping all relationship IDs."""
    # Collect all rIds from elements
    combined_xml = ' '.join(etree.tostring(e).decode() for e in elements_to_copy)
    flat_rids = set()
    for pair in re.findall(r'r:embed="(rId\d+)"|r:id="(rId\d+)"', combined_xml):
        for rid in pair:
            if rid: flat_rids.add(rid)

    rId_map = {}
    for rId in flat_rids:
        if rId not in input_doc.part.rels:
            continue
        rel = input_doc.part.rels[rId]
        try:
            new_rId = output_doc.part.relate_to(rel.target_part, rel.reltype)
            rId_map[rId] = new_rId
        except Exception:
            pass

    r_embed = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    r_id_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'

    for elem in elements_to_copy:
        elem_copy = deepcopy(elem)
        if rId_map:
            for node in elem_copy.iter():
                for attr in [r_embed, r_id_attr]:
                    old = node.get(attr)
                    if old and old in rId_map:
                        node.set(attr, rId_map[old])
        parent.insert(insert_pos, elem_copy)
        insert_pos += 1
    return insert_pos


def copy_sop_content(output_doc, input_doc, input_path):
    """
    Copy SOP content into SOP section:
    - If input has SOP heading → copy only that section's content
    - If input has no SOP heading → copy entire document
    - Always add [Please attach original SOP here] note first
    """
    token = '{{sop}}'
    placeholder_para = None
    for para in output_doc.paragraphs:
        if token in para.text:
            placeholder_para = para
            break
    if not placeholder_para:
        return

    # Clear placeholder
    for run in placeholder_para.runs:
        run.text = run.text.replace(token, '')

    parent = placeholder_para._element.getparent()
    idx = list(parent).index(placeholder_para._element)
    insert_pos = idx + 1

    # Add reference note first
    note = _make_note_para(
        '[ Please attach the original SOP / input MOP document here ]',
        color='C00000'
    )
    parent.insert(insert_pos, note)
    insert_pos += 1

    # Divider line
    divider = _make_note_para('─' * 60, color='808080')
    parent.insert(insert_pos, divider)
    insert_pos += 1

    # ── TXT file ──
    if input_doc is None:
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                txt_content = f.read().strip()
            if txt_content:
                txt_para = OxmlElement('w:p')
                txt_r = OxmlElement('w:r')
                txt_t = OxmlElement('w:t')
                txt_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                txt_t.text = txt_content
                txt_r.append(txt_t)
                txt_para.append(txt_r)
                parent.insert(insert_pos, txt_para)
        except Exception:
            pass
        return

    # ── DOCX file ──
    input_body = input_doc.element.body
    body_children = list(input_body)

    # Check if input has a recognizable SOP heading
    sop_start_idx = None
    sop_end_idx   = None
    para_idx_in_body = []

    # Map body children to paragraph objects for heading detection
    input_paras = input_doc.paragraphs
    para_elements = [p._element for p in input_paras]

    for body_i, child in enumerate(body_children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'p':
            continue
        if child not in para_elements:
            continue
        para_obj = input_paras[para_elements.index(child)]
        if 'heading' in para_obj.style.name.lower():
            matched = match_heading(para_obj.text.strip())
            if matched == 'sop':
                sop_start_idx = body_i + 1  # content starts after heading
            elif sop_start_idx is not None and sop_end_idx is None:
                sop_end_idx = body_i  # next heading = end of SOP section

    if sop_start_idx is not None:
        # Copy only SOP section content
        end = sop_end_idx if sop_end_idx else len(body_children)
        elements_to_copy = list(body_children[sop_start_idx:end])
        # Filter junk text-only paragraphs, keep tables and elements with images/OLE
        filtered = []
        for elem in elements_to_copy:
            tag = (elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag)
            if tag not in ('p', 'tbl', 'sdt'):
                continue
            elem_xml = etree.tostring(elem).decode()
            txt = ''.join(t for t in re.findall(r'<w:t[^>]*>([^<]*)</w:t>', elem_xml))
            has_rich = 'blip' in elem_xml or 'OLEObject' in elem_xml
            if has_rich or (txt.strip() and not is_junk(txt)):
                filtered.append(elem)
        if filtered:
            insert_pos = _copy_elements_with_rels(filtered, input_doc, output_doc, parent, insert_pos)
        else:
            # SOP section was empty — copy whole doc
            all_elems = [c for c in body_children if (c.tag.split('}')[-1] if '}' in c.tag else c.tag) in ('p', 'tbl', 'sdt')]
            insert_pos = _copy_elements_with_rels(all_elems, input_doc, output_doc, parent, insert_pos)
    else:
        # No SOP heading found — copy entire document
        all_elems = [c for c in body_children if (c.tag.split('}')[-1] if '}' in c.tag else c.tag) in ('p', 'tbl', 'sdt')]
        insert_pos = _copy_elements_with_rels(all_elems, input_doc, output_doc, parent, insert_pos)


def _update_header_date(docx_path):
    today = date.today().strftime('%Y-%m-%d')
    tmp_path = docx_path + '.hdrtmp'
    date_pattern = re.compile(r'\d{4}-\d{2}-\d{2}')
    try:
        with zipfile.ZipFile(docx_path, 'r') as zin:
            with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                seen = set()
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename in ['word/header1.xml', 'word/header2.xml']:
                        text = data.decode('utf-8')
                        text = date_pattern.sub(today, text)
                        data = text.encode('utf-8')
                    if item.filename not in seen:
                        zout.writestr(item, data)
                        seen.add(item.filename)
        os.replace(tmp_path, docx_path)
    except Exception:
        if os.path.exists(tmp_path):
            try: os.remove(tmp_path)
            except: pass


def generate_mop(template_path, input_path, activity_name, vendor_name, output_path):
    ext = os.path.splitext(input_path)[1].lower()
    input_doc = None
    if ext == '.docx':
        content_map, input_doc = extract_content_from_docx(input_path)
    elif ext == '.txt':
        content_map = extract_content_from_txt(input_path)
    else:
        return {'success': False, 'message': f'Unsupported format: {ext}. Upload .docx or .txt'}

    output_doc = Document(template_path)

    today = date.today().strftime("%d-%b-%Y")
    simple_replacements = {
        'version': '1.0', 'revdate': today,
        'prepared': 'Automation SME', 'change': 'Initial Release',
    }

    filled_sections, autofilled_sections = [], []
    all_keys = list(HEADING_MAP.keys())

    for key in all_keys:
        if key == 'sop':
            continue
        if key in content_map:
            replace_content_for_key(output_doc, input_doc, key, content_map[key], activity_name, vendor_name)
            filled_sections.append(key)
        else:
            auto_text = AUTO_FILL.get(key, f'{key} details for {activity_name} ({vendor_name})')
            auto_text = auto_text.replace('{activity}', activity_name).replace('{vendor}', vendor_name)
            token = '{{' + key + '}}'
            for para in output_doc.paragraphs:
                if token in para.text:
                    replace_placeholder_in_paragraph(para, key, auto_text)
                    break
            for table in output_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if token in para.text:
                                replace_placeholder_in_paragraph(para, key, auto_text)
            autofilled_sections.append(key)

    # SOP — copy ALL content from input file
    copy_sop_content(output_doc, input_doc, input_path)

    # Revision table
    replace_placeholders_in_doc(output_doc, simple_replacements)

    output_doc.save(output_path)
    _update_header_date(output_path)

    return {
        'success': True, 'message': 'MOP generated successfully.',
        'filled_sections': filled_sections,
        'autofilled_sections': autofilled_sections,
        'total_sections': len(all_keys)
    }

def sanitize_filename(name):
    name = re.sub(r'[^\w\s-]', '', name)
    name = re.sub(r'\s+', '_', name.strip())
    return name
