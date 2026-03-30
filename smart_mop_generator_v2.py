"""
Smart MOP Generator v2
======================
Upload Solution Document → Get Final MOP in exact template format.
No data is stored. All processing is in-memory only.

Install: pip install streamlit python-docx lxml
Run:     streamlit run smart_mop_generator.py
"""

import io
import re
import time
import copy
import shutil
import zipfile
import os
import tempfile
from datetime import datetime

import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ─────────────────────────────────────────────────────────────────
# PAGE CONFIG & CSS
# ─────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Smart MOP Generator",
    page_icon="📋",
    layout="centered",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:ital,wght@0,300;0,400;0,500;1,400&display=swap');

html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; }
.block-container { max-width: 800px; padding-top: 1.5rem; }

.hero {
    background: linear-gradient(135deg, #0d1b3e 0%, #0f2d4a 60%, #0a2218 100%);
    border: 1px solid rgba(99,179,237,0.18);
    border-radius: 18px;
    padding: 2.2rem 2rem 1.8rem;
    margin-bottom: 1.8rem;
    position: relative; overflow: hidden;
}
.hero::before {
    content:'';position:absolute;inset:0;
    background: radial-gradient(ellipse at 20% 50%, rgba(56,178,172,.07) 0%,transparent 60%),
                radial-gradient(ellipse at 80% 30%, rgba(99,179,237,.05) 0%,transparent 60%);
    pointer-events:none;
}
.hero-title {
    font-family:'Syne',sans-serif;font-size:2rem;font-weight:800;
    color:#e2e8f0;margin:0 0 .25rem;letter-spacing:-.5px;
}
.hero-title span{color:#63b3ed;}
.hero-sub{font-size:.88rem;color:#718096;margin:0;}
.badges{display:flex;gap:8px;margin-top:1.1rem;flex-wrap:wrap;}
.badge{font-size:.7rem;font-weight:500;padding:3px 10px;border-radius:20px;letter-spacing:.3px;}
.badge-g{background:rgba(56,178,172,.14);color:#38b2ac;border:1px solid rgba(56,178,172,.3);}
.badge-b{background:rgba(99,179,237,.12);color:#63b3ed;border:1px solid rgba(99,179,237,.25);}
.badge-o{background:rgba(237,137,54,.12);color:#ed8936;border:1px solid rgba(237,137,54,.25);}

.card{background:#111827;border:1px solid rgba(255,255,255,.07);border-radius:12px;
      padding:1.3rem 1.5rem;margin-bottom:1rem;}
.card h3{font-family:'Syne',sans-serif;font-size:.78rem;font-weight:700;color:#63b3ed;
          letter-spacing:1.4px;text-transform:uppercase;margin:0 0 .8rem;}

.privacy{background:rgba(56,178,172,.06);border-left:3px solid #38b2ac;
          border-radius:0 8px 8px 0;padding:.7rem 1rem;font-size:.78rem;
          color:#68d391;margin-bottom:1.4rem;}
.privacy strong{color:#9ae6b4;}

.pill-info{display:inline-flex;align-items:center;gap:6px;
            background:rgba(99,179,237,.1);border:1px solid rgba(99,179,237,.2);
            border-radius:6px;padding:5px 11px;font-size:.78rem;color:#90cdf4;margin:3px 0;}
.pill-ok{display:inline-flex;align-items:center;gap:6px;
          background:rgba(56,178,172,.1);border:1px solid rgba(56,178,172,.2);
          border-radius:6px;padding:5px 11px;font-size:.78rem;color:#81e6d9;margin:3px 0;}

.stButton>button{
    background:linear-gradient(135deg,#2b6cb0,#2c7a7b)!important;
    color:#fff!important;border:none!important;border-radius:10px!important;
    font-family:'Syne',sans-serif!important;font-weight:700!important;
    font-size:.95rem!important;padding:.65rem 2rem!important;width:100%!important;
    transition:all .2s!important;
}
.stButton>button:hover{
    background:linear-gradient(135deg,#3182ce,#319795)!important;
    transform:translateY(-1px)!important;
    box-shadow:0 8px 25px rgba(99,179,237,.22)!important;
}
[data-testid="stDownloadButton"]>button{
    background:linear-gradient(135deg,#276749,#285e61)!important;
    color:#fff!important;border:none!important;border-radius:10px!important;
    font-family:'Syne',sans-serif!important;font-weight:700!important;
    font-size:.95rem!important;padding:.65rem 2rem!important;width:100%!important;
}

.ps{display:flex;align-items:center;gap:10px;padding:7px 0;
     font-size:.8rem;border-bottom:1px solid rgba(255,255,255,.04);}
.ps.done{color:#68d391;} .ps.doing{color:#63b3ed;} .ps.wait{color:#4a5568;}
.pd{width:8px;height:8px;border-radius:50%;flex-shrink:0;}
.pd.done{background:#68d391;} .pd.doing{background:#63b3ed;animation:blink 1s infinite;}
.pd.wait{background:#2d3748;}
@keyframes blink{0%,100%{opacity:1}50%{opacity:.35}}
hr{border-color:rgba(255,255,255,.06)!important;}
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────
SECTION_KEYS = [
    "objective", "activity_description", "activity_type", "domain_in_scope",
    "prerequisites", "inventory_details", "node_connectivity", "iam",
    "triggering_method", "sop", "acceptance_criteria", "assumptions",
    "connectivity_diagram",
]

SECTION_LABELS = {
    "objective":            "1. Objective",
    "activity_description": "2. Activity Description",
    "activity_type":        "3. Activity Type",
    "domain_in_scope":      "4. Domain in Scope",
    "prerequisites":        "5. Pre-requisites",
    "inventory_details":    "6. Inventory Details",
    "node_connectivity":    "7. Node Connectivity Process",
    "iam":                  "8. Identity & Access Management",
    "triggering_method":    "9. Activity Triggering Method",
    "sop":                  "10. Standard Operating Procedure",
    "acceptance_criteria":  "11. Acceptance Criteria",
    "assumptions":          "12. Assumptions",
    "connectivity_diagram": "Connectivity Diagram:",
}

HEADING_MAP = {
    "objective":            ["objective"],
    "activity_description": ["activity description"],
    "activity_type":        ["activity type"],
    "domain_in_scope":      ["domain in scope", "domain"],
    "prerequisites":        ["pre-requisites", "prerequisites"],
    "inventory_details":    ["inventory details", "inventory"],
    "node_connectivity":    ["node connectivity", "node connectivity process"],
    "iam":                  ["identity & access", "identity and access", "identity"],
    "triggering_method":    ["activity triggering", "triggering method"],
    "sop":                  ["standard operating procedure", "sop"],
    "acceptance_criteria":  ["acceptance criteria"],
    "assumptions":          ["assumptions"],
    "connectivity_diagram": ["connectivity diagram"],
}

# TOC page numbers matching Final MOP output
TOC_PAGES = {
    "objective": 2, "activity_description": 2, "activity_type": 2,
    "domain_in_scope": 2, "prerequisites": 2, "inventory_details": 3,
    "node_connectivity": 3, "iam": 3, "triggering_method": 3,
    "sop": 3, "acceptance_criteria": 4, "assumptions": 4,
}

# Which sections use paragraph vs bullet vs numbered
PARA_SECTIONS   = {"objective", "activity_description", "activity_type",
                   "domain_in_scope", "inventory_details", "assumptions"}
BULLET_SECTIONS = {"prerequisites", "node_connectivity", "iam",
                   "triggering_method", "acceptance_criteria"}
NUMBERED_SECTIONS = {"sop"}


# ─────────────────────────────────────────────────────────────────
# DOCUMENT PARSING
# ─────────────────────────────────────────────────────────────────
def normalize_heading(text: str) -> str | None:
    t = re.sub(r'^\d+[\.\)]\s*', '', text).strip().lower()
    t = re.sub(r'\s+', ' ', t)
    for key, aliases in HEADING_MAP.items():
        for alias in aliases:
            if alias in t:
                return key
    return None


def extract_activity_name(doc: Document) -> str:
    """Get title from Heading1 or italic+underline near top."""
    for para in doc.paragraphs[:5]:
        if para.style.name.startswith("Heading 1"):
            txt = para.text.strip()
            # Remove "MOP:" prefix if present
            txt = re.sub(r'^MOP:\s*', '', txt, flags=re.IGNORECASE)
            return txt
    for para in doc.paragraphs[:8]:
        runs = para.runs
        if runs and runs[0].italic and runs[0].underline:
            return para.text.strip()
    return "Activity Name"


def extract_sections(doc: Document) -> dict:
    """
    Parse solution document.
    Returns dict: key -> list of paragraphs (strings).
    Images are stored as list of (rId, image_bytes, ext) tuples under key 'connectivity_diagram'.
    """
    sections = {k: [] for k in SECTION_KEYS}
    sections["connectivity_diagram"] = []  # will hold image tuples
    current_key = None
    image_rels = {}  # rId -> (bytes, ext)

    # Build image relationship map from document
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                ext = rel.target_part.content_type.split("/")[-1]
                if ext == "jpeg":
                    ext = "jpg"
                image_rels[rel.rId] = (img_bytes, ext)
            except Exception:
                pass

    for para in doc.paragraphs:
        style = para.style.name
        text  = para.text.strip()

        # Detect heading
        if style.startswith("Heading"):
            key = normalize_heading(text)
            if key:
                current_key = key
            continue

        if current_key is None:
            continue

        # Skip TOC / cover lines
        if re.match(r'^\d+\.\s+\w.*Page\s+\d+', text):
            continue
        # Check for inline images FIRST (image paras may have empty text)
        has_image = False
        _BLIP = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        _REL  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        for blip in para._p.findall(f'.//{{{_BLIP}}}blip'):
            embed = blip.get(f'{{{_REL}}}embed')
            if embed and embed in image_rels:
                sections["connectivity_diagram"].append(image_rels[embed])
                has_image = True

        if has_image:
            continue

        if text in ("METHOD OF PROCEDURE", "CONTENTS:", "CONTENTS", ""):
            continue

        if text and current_key in sections:
            # Clean leading dashes/numbers for list items
            clean = re.sub(r'^[-–•]\s*', '', text)
            clean = re.sub(r'^\d+[\.\)]\s*', '', clean)
            sections[current_key].append(clean.strip())

    return sections


# ─────────────────────────────────────────────────────────────────
# DOCX BUILDER HELPERS
# ─────────────────────────────────────────────────────────────────
def add_run(para, text, font="Calibri", size=None, bold=False,
            italic=False, underline=False, color=None):
    run = para.add_run(text)
    run.font.name = font
    if size:
        run.font.size = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def set_right_tab(para, pos=8400):
    pPr = para._p.get_or_add_pPr()
    # Remove existing tabs
    for old in pPr.findall(qn("w:tabs")):
        pPr.remove(old)
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos))
    tabs.append(tab)
    pPr.append(tabs)


def page_break(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_header(doc, today_str: str):
    """4-line header exactly matching Final MOP."""
    section = doc.sections[0]
    header  = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    # Line 1 — empty (spacing)
    header.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Line 2 — ERICSSON bold blue
    p2 = header.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, "ERICSSON", size=22, bold=True, color=(0, 58, 143))

    # Line 3 — Ericsson Confidential [tab] Document Type
    p3 = header.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    set_right_tab(p3, 9000)
    add_run(p3, "Ericsson Confidential", size=11)
    add_run(p3, "\t", size=11)
    add_run(p3, "Document Type: Requirement Specification", size=11)

    # Line 4 — Prepared By ... Approved By
    p4 = header.add_paragraph()
    p4.paragraph_format.space_after = Pt(0)
    add_run(p4, "Prepared By: Automation SME", size=11)
    add_run(p4, "                                           ", size=11)
    add_run(p4, "Approved By: ____________", size=11)

    # Line 5 — Date (auto today)
    p5 = header.add_paragraph()
    p5.paragraph_format.space_after = Pt(0)
    add_run(p5, f"Date: {today_str}", size=11)

    # Blank line
    header.add_paragraph()


def add_footer(doc):
    """Center-aligned  Page [N]  footer."""
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp, "Page ", size=11)

    fldChar1   = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText  = OxmlElement("w:instrText"); instrText.text = "PAGE"
    fldChar2   = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3   = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    run = fp.add_run()
    run.font.name = "Calibri"; run.font.size = Pt(11)
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])


def add_heading2(doc, text: str):
    """Heading 2 style: Calibri 13pt, bold, #4F81BD, space_before=10pt."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(13)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)


def add_body_para(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, size=11)


def add_bullet_item(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, size=11)


def add_numbered_item(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, size=11)


def add_image_to_doc(doc, img_bytes: bytes, ext: str):
    """Insert image into document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    img_stream = io.BytesIO(img_bytes)
    # Width = 5 inches (matching original cx=2508250 EMU ≈ 2.74 inches, keep proportional)
    run.add_picture(img_stream, width=Inches(5))


# ─────────────────────────────────────────────────────────────────
# MAIN MOP BUILDER
# ─────────────────────────────────────────────────────────────────
def build_mop(activity_name: str, sections: dict, today_str: str) -> bytes:
    doc = Document()

    # ── Page setup: US Letter, matching Final MOP margins ─────────
    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1.25)   # 1800 DXA
    sec.right_margin  = Inches(1.25)
    sec.header_distance = Inches(0.5)
    sec.footer_distance = Inches(0.5)

    add_header(doc, today_str)
    add_footer(doc)

    # ══════════════════════════════════════════════════════════════
    # PAGE 1 — Cover Page
    # ══════════════════════════════════════════════════════════════

    # "METHOD OF PROCEDURE" — center, bold, gray #7F7F7F, 18pt
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(4)
    add_run(tp, "METHOD OF PROCEDURE", size=18, bold=True, color=(0x7F, 0x7F, 0x7F))

    # Activity Name — indent left=1440 firstLine=720, italic, underline, 14pt
    ap = doc.add_paragraph()
    ap.paragraph_format.space_after = Pt(4)
    pPr = ap._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "1440"); ind.set(qn("w:firstLine"), "720")
    pPr.append(ind)
    add_run(ap, activity_name, size=14, italic=True, underline=True)

    # Blank
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # "CONTENTS:" — bold, underline, 12pt
    cp = doc.add_paragraph()
    cp.paragraph_format.space_after = Pt(4)
    add_run(cp, "CONTENTS:", size=12, bold=True, underline=True)

    # TOC entries — right tab at 8400
    for key in SECTION_KEYS[:-1]:   # skip connectivity_diagram
        label    = SECTION_LABELS[key]
        page_num = TOC_PAGES.get(key, 2)
        tp2 = doc.add_paragraph()
        tp2.paragraph_format.space_after = Pt(2)
        set_right_tab(tp2, 8400)
        add_run(tp2, label, size=11)
        add_run(tp2, f"\tPage {page_num}", size=11)

    # Page break → Section content starts on page 2
    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # SECTIONS 1–12  (all on page 2 onward, no forced page breaks)
    # ══════════════════════════════════════════════════════════════
    for key in SECTION_KEYS[:-1]:   # skip connectivity_diagram
        label   = SECTION_LABELS[key]
        content = sections.get(key, [])

        add_heading2(doc, label)

        if key in PARA_SECTIONS:
            text = " ".join(content).strip()
            if text:
                add_body_para(doc, text)
            else:
                add_body_para(doc, "")

        elif key in BULLET_SECTIONS:
            for item in content:
                add_bullet_item(doc, item)

        elif key in NUMBERED_SECTIONS:
            for item in content:
                add_numbered_item(doc, item)

    # ── Connectivity Diagram section ──────────────────────────────
    images = sections.get("connectivity_diagram", [])
    if images:
        add_heading2(doc, SECTION_LABELS["connectivity_diagram"])
        for img_bytes, ext in images:
            add_image_to_doc(doc, img_bytes, ext)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────
# STREAMLIT UI
# ─────────────────────────────────────────────────────────────────

# Hero
st.markdown("""
<div class="hero">
  <p class="hero-title">Smart <span>MOP</span> Generator</p>
  <p class="hero-sub">Upload your Solution Document → Get a fully formatted MOP instantly</p>
  <div class="badges">
    <span class="badge badge-g">⚡ In-Memory Only</span>
    <span class="badge badge-b">📋 Auto TOC + Page Numbers</span>
    <span class="badge badge-o">🖼️ Images Preserved</span>
  </div>
</div>
""", unsafe_allow_html=True)

# Privacy
st.markdown("""
<div class="privacy">
  <strong>🔒 Zero Data Storage:</strong> Your files are processed entirely in-memory.
  Nothing is written to disk or any database. All data is discarded when your session ends.
</div>
""", unsafe_allow_html=True)

# ── Upload ────────────────────────────────────────────────────────
st.markdown('<div class="card"><h3>📂 Upload Solution Document</h3>', unsafe_allow_html=True)
sol_file = st.file_uploader(
    "Upload Solution Document (.docx)",
    type=["docx"],
    label_visibility="collapsed",
    key="sol_upload"
)
if sol_file:
    st.markdown(f'<div class="pill-ok">✅ Loaded: <strong>{sol_file.name}</strong> &nbsp;·&nbsp; {sol_file.size/1024:.1f} KB</div>',
                unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)

# ── Generate ──────────────────────────────────────────────────────
st.markdown('<div class="card"><h3>⚡ Generate MOP</h3>', unsafe_allow_html=True)
gen_btn = st.button("🚀 Generate MOP Document", disabled=not sol_file)
st.markdown('</div>', unsafe_allow_html=True)

if gen_btn and sol_file:

    st.markdown('<div class="card"><h3>⚙️ Processing</h3>', unsafe_allow_html=True)

    steps = [
        "Reading solution document",
        "Extracting activity name",
        "Parsing all 12 sections",
        "Detecting images & diagrams",
        "Building cover page & TOC",
        "Populating section content",
        "Applying header & footer",
        "Finalising MOP document",
    ]

    placeholders = [st.empty() for _ in steps]
    for ph, s in zip(placeholders, steps):
        ph.markdown(f'<div class="ps wait"><div class="pd wait"></div>{s}</div>', unsafe_allow_html=True)

    try:
        activity_name = ""; sections = {}; today_str = ""
        output_bytes  = b""

        for i, (ph, step) in enumerate(zip(placeholders, steps)):
            ph.markdown(f'<div class="ps doing"><div class="pd doing"></div>{step}</div>', unsafe_allow_html=True)
            time.sleep(0.2)

            if i == 0:
                sol_bytes = sol_file.read()
                sol_doc   = Document(io.BytesIO(sol_bytes))
            elif i == 1:
                activity_name = extract_activity_name(sol_doc)
                today_str     = datetime.today().strftime("%d %B %Y")
            elif i == 2:
                sections = extract_sections(sol_doc)
            elif i == 3:
                n_images = len(sections.get("connectivity_diagram", []))
            elif i == 7:
                output_bytes = build_mop(activity_name, sections, today_str)

            ph.markdown(f'<div class="ps done"><div class="pd done"></div>{step} ✓</div>', unsafe_allow_html=True)
            time.sleep(0.05)

        st.markdown('</div>', unsafe_allow_html=True)

        # ── Success ───────────────────────────────────────────────
        st.markdown(f"""
        <div style="background:rgba(56,178,172,.08);border:1px solid rgba(56,178,172,.25);
             border-radius:12px;padding:1.4rem;margin:1rem 0;text-align:center;">
          <div style="font-family:'Syne',sans-serif;font-size:1.05rem;font-weight:700;
               color:#9ae6b4;margin-bottom:.3rem;">✅ MOP Generated Successfully</div>
          <div style="font-size:.8rem;color:#68d391;">
            Activity: <strong style="color:#9ae6b4;">{activity_name}</strong>
          </div>
        </div>""", unsafe_allow_html=True)

        safe = re.sub(r'[^\w\s-]', '', activity_name).strip().replace(' ', '_')[:60]
        st.download_button(
            label="📥 Download MOP.docx",
            data=output_bytes,
            file_name=f"MOP_{safe}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # Summary
        st.markdown('<div class="card"><h3>📊 Content Summary</h3>', unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        filled = sum(1 for k in SECTION_KEYS[:-1] if sections.get(k))
        with c1: st.metric("Sections Filled", f"{filled}/12")
        with c2: st.metric("Images Found", len(sections.get("connectivity_diagram", [])))
        with c3:
            total = sum(len(v) for k, v in sections.items() if k != "connectivity_diagram")
            st.metric("Content Lines", total)
        st.markdown('</div>', unsafe_allow_html=True)

        # Preview
        with st.expander("📋 Preview extracted content"):
            for key in SECTION_KEYS[:-1]:
                content = sections.get(key, [])
                label   = SECTION_LABELS[key]
                if content:
                    st.markdown(f"**{label}**")
                    for line in content[:3]:
                        st.markdown(f"<span style='color:#a0aec0;font-size:.78rem;'>• {line[:130]}</span>",
                                    unsafe_allow_html=True)
                    if len(content) > 3:
                        st.caption(f"... +{len(content)-3} more")
                else:
                    st.markdown(f"<span style='color:#4a5568;font-size:.78rem;'>{label} — empty</span>",
                                unsafe_allow_html=True)

    except Exception as e:
        st.markdown('</div>', unsafe_allow_html=True)
        st.error(f"❌ Error: {e}")
        import traceback
        st.code(traceback.format_exc())

elif gen_btn:
    st.warning("⚠️ Please upload a Solution Document first.")

# Footer
st.markdown("""<br>
<div style="text-align:center;font-size:.7rem;color:#2d3748;padding:.8rem 0;
     border-top:1px solid rgba(255,255,255,.04);">
  🔒 No data stored &nbsp;·&nbsp; In-memory processing only &nbsp;·&nbsp;
  Session cleared on close &nbsp;·&nbsp; Smart MOP Generator v2
</div>""", unsafe_allow_html=True)
